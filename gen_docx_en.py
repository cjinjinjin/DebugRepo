#!/usr/bin/env python3
"""Generate DLIS_Model_Deployment_Guide_v5_EN.docx — English version."""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0, 51, 102)

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    return p

def add_link_note(text):
    p = doc.add_paragraph()
    run = p.add_run('-> ' + text)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.font.size = Pt(10)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    shd = run._element.makeelement(qn('w:shd'), {})
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F0F0F0')
    rpr = run._element.find(qn('w:rPr'))
    if rpr is None:
        rpr = run._element.makeelement(qn('w:rPr'), {})
        run._element.insert(0, rpr)
    rpr.append(shd)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri+1].cells[ci].text = str(val)
    doc.add_paragraph()

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

# -- Image paths --
IMG_DIR = r'C:\Users\jinjinchen\OneDrive - Microsoft\doc_images'

def add_image(filename, caption=None, width=Inches(5.5)):
    import os
    img_path = os.path.join(IMG_DIR, filename)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=width)
        if caption:
            p = doc.add_paragraph()
            run = p.add_run(caption)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(120, 120, 120)
    else:
        add_note(f'[Image missing: {filename}]')

# -- Document links --
XUCHA_DOC_URL = 'https://microsoftapc-my.sharepoint.com/:w:/g/personal/xucha_microsoft_com/IQCgQ8ufJRUWSoedwJNRll1yAc-aqBPPV5Rqmg4JjANvS2E'
ZHAHAO_DOC_URL = 'https://microsoftapc-my.sharepoint.com/:w:/g/personal/zhahao_microsoft_com/IQDV9SZPUx2jRbIULT1vcKGTAXmrH1bD-lcaqsz2L80tl_0'
CHUNCHEN_DOC_URL = 'https://microsoftapc-my.sharepoint.com/:w:/g/personal/chunchen_microsoft_com/IQDG4fXi2Zd3RqUUqu9dTNLqAaeNMD_4epy-2DcbxOt1pq8'

# -- External resource URLs --
DLIS_WIKI_LLM_ENGINE = 'https://dlisinfrawiki.azurewebsites.net/wiki/Contents/DLISLLMEngine/Optimizing+a+LLM+model+with+DLIS+LLM+engine.html'
DLIS_WIKI_DATA_TRANSFER = 'https://dlisinfrawiki.azurewebsites.net/wiki/Contents/ModelRepository/Data+Transfer+tools.html'
DLIS_WIKI_CENTRAL_LOG = 'https://dlisinfrawiki.azurewebsites.net/wiki/Contents/Debugging/How+to+use+Central+Log.html'
ONE_INFERENCE_PORTAL = 'https://dlis-portal.microsoft.com/'
OAAS_LLM_TEMPLATE_REPO = 'https://msasg.visualstudio.com/Bing_and_IPG/_git/OaaS_LLMTemplate'
OAAS_LLM_TEMPLATE_PIPELINES = 'https://msasg.visualstudio.com/Bing_and_IPG/_build?definitionScope=%5COaaS_LLMTemplate'
DLIS_COPY_PIPELINE = 'https://msasg.visualstudio.com/Bing_and_IPG/_build?definitionId=43927'
COSMOS_REPO_URL = 'https://cosmos09.osdinfra.net:443/cosmos/DLISModelRepository/'
KUSTO_SI_URL = 'https://bingadsppe.kusto.windows.net/'
KUSTO_PROD_URL = 'https://bingads.kusto.windows.net/'
JARVIS_PROD_URL = 'https://jarvis-west.dc.ad.msft.net/dashboard/DLIS-Model-Metrics'
JARVIS_SI_URL = 'https://jarvis-west.dc.ad.msft.net/dashboard/DLIS-Model-Metrics-SI'
SIWEN_DESHENG_RECORDING = 'https://teams.microsoft.com/l/meetingrecap?driveId=b%21FSzYuSuEp02FCYr9pdB1t_R60kGlONJJht32xiFgk6Df_y20EhFwRaIwLXCzCWLd&driveItemId=015S6O6HUNTJFSKRTY3NH24GGIFES2P7NK&sitePath=https%3A%2F%2Fmicrosoftapc-my.sharepoint.com%2Fpersonal%2Fsiwenzhu_microsoft_com%2FDocuments%2FRecordings%2FCall+with+Desheng+Cui-20260326_203633-Meeting+Recording.mp4%3Fweb%3D1&fileUrl=https%3A%2F%2Fmicrosoftapc-my.sharepoint.com%2Fpersonal%2Fsiwenzhu_microsoft_com%2FDocuments%2FRecordings%2FCall+with+Desheng+Cui-20260326_203633-Meeting+Recording.mp4%3Fweb%3D1&threadId=19%3A01d7b57a-2d8a-45ef-b6eb-9d9cd1c21105_0b99169c-642e-4bbd-afec-02523f1445b5%40unq.gbl.spaces&callId=a347d990-0384-4b69-a3a9-00885914510e&threadType=OneOnOneChat&meetingType=Unknown&subType=RecapSharingLink_RecapCore'

def add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')
    rPr.append(sz)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph

def add_para_with_link(prefix, link_text, url, suffix=''):
    """Add a paragraph with inline hyperlink."""
    p = doc.add_paragraph()
    if prefix:
        p.add_run(prefix)
    add_hyperlink(p, link_text, url)
    if suffix:
        p.add_run(suffix)
    return p

def add_note_with_link(prefix, link_text, url, suffix=''):
    """Add a note-styled paragraph with inline hyperlink."""
    p = doc.add_paragraph()
    if prefix:
        run = p.add_run(prefix)
        run.italic = True
        run.font.color.rgb = RGBColor(0, 102, 153)
    add_hyperlink(p, link_text, url)
    if suffix:
        run = p.add_run(suffix)
        run.italic = True
        run.font.color.rgb = RGBColor(0, 102, 153)
    return p

def add_table_with_links(headers, rows):
    """Add a table where cell values can be (text, url) tuples for hyperlinks."""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            if isinstance(val, tuple):
                text, url = val
                cell.text = ''
                p = cell.paragraphs[0]
                add_hyperlink(p, text, url)
            else:
                cell.text = str(val)
    doc.add_paragraph()

def add_doc_link(label, url):
    """Add a document reference link with clickable hyperlink."""
    p = doc.add_paragraph()
    run = p.add_run('[Attachment] ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.font.size = Pt(10)
    add_hyperlink(p, label, url)

# ================================================================
# TITLE
# ================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DLIS Model Deployment Guide (OaaS_LLMTemplate)')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Version v5.2 \u2014 2026-04-23\nConsolidating Gemma4, ZImage deployment experience, ChangXu v2 doc, and team insights from Teams chats')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ================================================================
# TOC
# ================================================================
add_heading('Table of Contents', 1)
toc_items = [
    'Part 1: Deployment Main Flow',
    '  1. Introduction',
    '  2. Deployment Flow Overview',
    '  3. Step 1: Local Development & Testing',
    '  4. Step 2: Upload Checkpoint to Gen1',
    '  5. Step 3: Gen1 to Gen2 Data Migration',
    '  6. Step 4: PR Submission & CI Auto Image Build',
    '  7. Step 5: Polaris Testing',
    '  8. Step 6: DLIS Production Deployment',
    '  9. Step 7: Post-Deployment Verification',
    '  10. Step 8: Polaris Job Optimization (Optional)',
    '',
    'Part 2: Environment & Operations',
    '  11. SI/Prod Environments & Certificate Management',
    '  12. Kusto Log Viewing & Debugging',
    '  13. Common Issues & Solutions',
    '',
    'Appendices',
    '  A. Base Docker Image Selection & Updates',
    '  B. Tsinghua Mirror Configuration',
    '  C. Model Code Writing Guide',
    '  D. External settings.json Configuration Override',
    '  E. Central Log Debugging Tool',
    '  F. Reference Links',
    '  G. Gemma4 Deployment Iteration Summary',
    '  H. Revision History',
]
for item in toc_items:
    if item == '':
        doc.add_paragraph()
    elif item.startswith('Part') or item.startswith('Appendices'):
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.bold = True
    else:
        doc.add_paragraph(item.strip(), style='List Bullet')

doc.add_page_break()

# ================================================================
# PART 1: MAIN FLOW
# ================================================================
p = doc.add_paragraph()
run = p.add_run('Part 1: Deployment Main Flow')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)

# ========== Section 1 ==========
add_heading('1. Introduction', 1)
add_heading('1.1 Purpose of This Document', 2)
add_para('This document provides a complete guide for deploying LLM / multimodal / Diffusion models to the DLIS (Deep Learning Inference Service) platform.')

add_heading('1.2 What is OaaS?', 2)
add_para('OaaS (Optimization as a Service) is a comprehensive framework for deploying LLMs on DLIS, providing:')
add_bullet('Automated model optimization: standardized workflows for quantization, compilation, and optimization')
add_bullet('Multi-backend support: vLLM and TensorRT-LLM inference engines')
add_bullet('Production-ready templates: pre-configured templates for rapid deployment')
add_bullet('Batch processing support: efficient batch inference for high-throughput scenarios')
add_note_with_link('Official docs: ', 'Optimizing a LLM model with DLIS LLM engine (DLIS Wiki)', DLIS_WIKI_LLM_ENGINE)

# ========== Section 2 ==========
add_heading('2. Deployment Flow Overview', 1)
add_code('''  ┌───────────────────────────────────┐
  │ Step 1: Local Dev & Docker Test    │  ← Verify correctness first
  └────────────────┬──────────────────┘
                   │
      ┌────────────┴────────────┐
      ▼                         ▼         (can run in parallel)
  ┌─────────────────────┐  ┌─────────────┐
  │Step 2: Upload ckpt  │  │Step 4:      │
  │     to Gen1         │  │PR & CI      │
  │        ↓            │  │Build Image  │
  │Step 3: Gen1→Gen2    │  │             │
  │     Migration       │  │             │
  └──────────┬──────────┘  └──────┬──────┘
             └────────────┬───────┘
                          ▼
  ┌───────────────────────────────────┐
  │ Step 5: Polaris Test               │  ← Check output & latency; DLIS reads model from Gen2
  └────────────────┬──────────────────┘
                   ▼
  ┌───────────────────────────────────┐
  │ Step 6: DLIS Production Deploy     │
  └────────────────┬──────────────────┘
                   ▼
  ┌───────────────────────────────────┐
  │ Step 7: Verification & Monitoring  │
  └───────────────────────────────────┘''')

add_para('Key Principles:', bold=True)
add_bullet('Test locally first. Only upload data and build images after local Docker tests pass')
add_bullet('Integrate Kusto logging from the local testing phase. Do not wait until online deployment — configure certificates and EventHub during local Docker testing so you can view logs in Kusto in real-time and catch log format, auth, and connection issues early')
add_bullet('DLIS reads model data from Gen2 (dlisstoregen2.dfs.core.windows.net), not Gen1')
add_bullet('Images are automatically built via the OaaS_LLMTemplate repo CI pipeline')
p = doc.paragraphs[-1]
p.clear()
p.add_run('Images are automatically built via the ')
add_hyperlink(p, 'OaaS_LLMTemplate repo', OAAS_LLM_TEMPLATE_REPO)
p.add_run(' CI pipeline')
add_bullet('Step 2+3 (upload ckpt + data migration) and Step 4 (image build) can run in parallel, independently')

# ========== Section 3 ==========
add_heading('3. Step 1: Local Development & Testing', 1)
add_heading('3.1 Development Flow Overview', 2)
add_para_with_link('Create a personal branch in the ', 'OaaS_LLMTemplate repo', OAAS_LLM_TEMPLATE_REPO, ' (e.g., jinjinchen/ZImage-v1). Complete the following before pushing code to trigger CI build:')
add_bullet('1. Modify model code (model.py, dlis_inter.py, etc.)')
add_bullet('2. Evaluate whether OaaS template customization is needed (multimodal support, custom formats, etc.)')
add_bullet('3. Choose the appropriate Dockerfile (fast iteration vs full build)')
add_bullet('4. Build and test locally with Docker, confirm correct functionality')

add_heading('3.2 Code File Reference', 2)
add_para_with_link('In the ', 'OaaS_LLMTemplate repo', OAAS_LLM_TEMPLATE_REPO, ', create a personal branch (e.g., jinjinchen/ZImage-v1) and modify the following files:')
add_table(
    ['File', 'Description'],
    [
        ['dlis_model/model/model.py', 'Model initialization + inference logic (core file)'],
        ['dlis_model/model/dlis_inter.py', 'Pre/post-processing, implements PreAndPostProcessor class'],
        ['dlis_model/http_server.py', 'HTTP server (if custom format needed)'],
        ['requirements-vllm.txt', 'Python dependencies'],
    ]
)
add_link_note('See "Appendix C: Model Code Writing Guide" for model.py details')

add_heading('3.3 OaaS Template Customization (Optional)', 2)
add_note_with_link('Template customization is optional. First check whether the original ', 'OaaS LLM Template repo', OAAS_LLM_TEMPLATE_REPO, ' template meets your needs. For pure text LLM inference, customization is usually unnecessary.')
add_table(
    ['Customization Need', 'Description', 'File to Modify'],
    [
        ['Multimodal vLLM support', 'Original template only supports text input', 'model.py'],
        ['Image transfer format', 'Use multipart/form-data for efficiency', 'http_server.py'],
        ['Additional dependencies', 'e.g., diffusers, Pillow, etc.', 'requirements-vllm.txt'],
        ['Non-LLM models', 'e.g., Diffusion models (ZImage)', 'model.py'],
    ]
)

add_heading('3.4 Two Dockerfile Options', 2)
add_table(
    ['Dockerfile', 'Base Image', 'Build Time', 'Use Case'],
    [
        ['Dockerfile_vllm_0.10.0', 'nvidia/cuda:12.8.1-devel-ubuntu22.04', 'Tens of minutes', 'Need specific vLLM/torch version'],
        ['Dockerfile_vllm_fast', 'vllm/vllm-openai:latest', '< 1 second', 'Fast iteration'],
    ]
)
add_link_note('See "Appendix A: Base Docker Image Selection & Updates" for version compatibility details')

add_para('Common Docker Build Network Issues:', bold=True)
add_bullet('apt-get install fails: CI build agent cannot connect to archive.ubuntu.com')
add_bullet('pip install timeout: pypi.org unreachable')
add_link_note('See "Appendix B: Tsinghua Mirror Configuration" for solutions')

add_heading('3.5 Local Docker Testing', 2)
add_code('''# 1. Build image
cd /path/to/OaaS_LLMTemplate
export SOURCE_BRANCH="test"
sudo bash pipeline/build_vllm_image.sh

# 2. Start container
IMAGE_TAG="<build_tag>"
sudo docker run -d --name model-test \\
  --gpus all \\
  -v /path/to/model_weights:/Model/model_name \\
  -p <host_port>:8888 \\
  <image_name>:$IMAGE_TAG \\
  /dlis_model/run.sh http

# 3. Test request
curl -X POST http://localhost:<host_port> \\
  -H "Content-Type: application/json" \\
  -d '{"prompt": "test input"}'
''')

add_para('Important Notes:', bold=True)
add_bullet('If a GPU is occupied, use --gpus \'"device=N"\' to specify an available GPU')
add_bullet('Must add -p <host_port>:8888 port mapping, otherwise host curl will get 403')
add_bullet('vllm/vllm-openai base image entrypoint is vllm serve; use --entrypoint bash for interactive shell')
add_bullet('If Python files are volume-mounted, __pycache__ may cause stale code to be loaded')

add_para('Verify Kusto Logging (recommended during local testing):', bold=True)
add_para('After starting the local Docker container, verify Kusto log delivery alongside inference testing. This catches certificate and EventHub configuration issues before deployment.')
add_bullet('1. Ensure the correct environment PFX certificate (e.g., AggSvcAuthCert-si.pfx) is in your Cosmos directory and volume-mounted to /Model in the container')
add_bullet('2. Ensure settings.json has the correct EventHub namespace and kusto_log parameters (cert path, topic, etc.)')
add_bullet('3. After sending a test request, check container logs (docker logs) for EventHub send success/failure output')
add_bullet('4. Query the corresponding Kusto environment table in Kusto Explorer to confirm logs arrived (typically 1-2 minute delay)')
add_code('''# SI environment Kusto query example
// Kusto cluster: https://bingadsppe.kusto.windows.net/
// Database: appsvc
appsvc_info
| where TIMESTAMP > ago(10m)
| where ModelName == "<your_model_name>"
| order by TIMESTAMP desc
| take 20''')
add_note('If no results: ① check cert environment matches namespace (SI cert + SI namespace); ② verify logger level is set to INFO; ③ check if EventHub send errors are silently swallowed. See Section 12 for details.')

add_heading('3.6 Offline Testing (No HTTP Server)', 2)
add_code('''sudo docker run --rm -it --gpus all \\
  -v /path/to/model:/Model/model_name \\
  <image>:<tag> \\
  bash -c 'cd /dlis_model && ./run.sh offline /tmp/input.json /tmp/output.json'
''')

# ========== Section 4 ==========
add_heading('4. Step 2: Upload Checkpoint to Gen1', 1)
add_para('Upload model weight files to Gen1 Cosmos storage.')
add_heading('4.1 Upload Destination', 2)
add_code('Target: https://cosmos09.osdinfra.net:443/cosmos/DLISModelRepository/local/<your-alias>/')
add_para('Use Visual Studio Scope Extension for authentication and upload.')
add_image('xucha_gen1_upload.png', 'Fig: Gen1 Cosmos upload directory example (Source: ChangXu doc)')
add_doc_link('Source doc: DLIS_Model_DeploymentWith_OaaS_v2.docx (ChangXu)', XUCHA_DOC_URL)

add_heading('4.2 Directory Structure (Flat Layout, Recommended)', 2)
add_code('''<model-dir>/
|-- model_name/               <- Model weight files folder
|   |-- config.json
|   |-- model-00001-of-N.safetensors
|   |-- tokenizer.json
|   +-- ...
|-- dlis_inter.py             <- Place in root directory
|-- settings.json             <- Config override (optional)
|-- AggSvcAuthCert-prod.pfx   <- Kusto certificate (optional)
+-- AggSvcAuthCert-si.pfx
''')

add_para('Critical Notes:', bold=True)
add_bullet('dlis_inter.py must be in the Cosmos root directory; sys.path.append(\'/Model\') only finds root-level files')
add_bullet('Do NOT put model.py in Cosmos; use the image-bundled version. Old model.py on Cosmos will override the new version in the image')
add_bullet('Do NOT upload unnecessary large files (e.g., .tar packages) \u2014 wastes sync time')
add_bullet('Place certificates in the Cosmos root directory, not in subdirectories (DLIS only mounts the top-level directory)')
add_link_note('See "Appendix D: External settings.json Configuration Override" for settings.json details')

# ========== Section 5 ==========
add_heading('5. Step 3: Gen1 to Gen2 Data Migration', 1)
add_note('Note: Step 2+3 (upload ckpt + data migration) and Step 4 (image build) can run in parallel, independently.')
add_para('DLIS reads model data from Gen2 at deployment time, so Gen1 data must be migrated to Gen2.')
add_para_with_link('Reference Wiki: ', 'Data Transfer Tools (DLIS Wiki)', DLIS_WIKI_DATA_TRANSFER)
add_para_with_link('Gen1 to Gen2 Migration Steps (see ', 'How_to_Build_Your_Own_DLIS_Model.docx Step 6.2', ZHAHAO_DOC_URL, '):')
add_bullet('1. Create a branch in Repos')
add_bullet('2. Open DLIS copy pipeline, select View/Edit')
p = doc.paragraphs[-1]
p.clear()
p.add_run('2. Open ')
add_hyperlink(p, 'DLIS copy pipeline', DLIS_COPY_PIPELINE)
p.add_run(', select View/Edit')
add_image('zhahao_image002.png', 'Fig: Create branch and open pipeline (Source: Hao Zhang doc)')
add_bullet('3. Select the newly created branch and update parameters')
add_image('zhahao_image004.png', 'Fig: Select branch and set pipeline variables (Source: Hao Zhang doc)')
add_image('zhahao_image006.png', 'Fig: Configure migration path parameters (Source: Hao Zhang doc)')
add_bullet('4. Click Validate and Save to save parameters')
add_image('zhahao_image008.png', 'Fig: Save and run pipeline (Source: Hao Zhang doc)')
add_doc_link('Source doc: How_to_Build_Your_Own_DLIS_Model.docx (Hao Zhang)', ZHAHAO_DOC_URL)
add_code('Gen2 path format:\nabfs://dlisstore@dlisstoregen2.dfs.core.windows.net/dlismodelrepository-c09/local/users/<username>/<model-dir>/')

add_para('Gen2 Verification (Important):', bold=True)
add_bullet('After migration, use SAW (Secure Admin Workstation) to verify file completeness on Gen2 (Desheng\'s advice)')
add_bullet('Files existing on Gen2 does not guarantee correctness \u2014 verify file sizes and integrity')
add_bullet('ADL data migration tools have pitfalls; carefully inspect after migration')

add_para('ModelDataPath Mechanism:', bold=True)
add_bullet('ModelDataPath points to a file on Gen2 (e.g., complete.txt); DLIS actually mounts the entire parent directory to /Model')
add_bullet('/Model is mounted read-only \u2014 cannot create files in it')
add_bullet('If runtime config files are needed, use the writable mirror approach (see Common Issues)')

# ========== Section 6 ==========
add_heading('6. Step 4: PR Submission & CI Auto Image Build', 1)
add_note('Note: Step 2+3 (upload ckpt + data migration) and Step 4 (image build) can run in parallel, independently.')
add_para_with_link('After local development and testing, push code to your personal branch in the ', 'OaaS_LLMTemplate repo', OAAS_LLM_TEMPLATE_REPO, '. The CI pipeline will automatically build the Docker image.')

add_heading('6.1 Create Branch and Push Code', 2)
add_code('''# Create a personal branch in OaaS_LLMTemplate repo
git checkout -b <your-alias>/<model-name>   # e.g., jinjinchen/ZImage-v1

# Commit local changes
git add -A
git commit -m "Add <model-name> model support"
git push origin <your-alias>/<model-name>
''')

add_heading('6.2 CI Pipeline Auto Build', 2)
add_para('After pushing to the branch, the CI pipeline automatically triggers an image build. No manual action required.')
add_table(
    ['Item', 'Description'],
    [
        ['Trigger', 'Push to any branch triggers automatically (including non-main branches)'],
        ['Image Tag Format', 'YYYYMMDD-HHMM-<branch_name> (non-main branches)'],
        ['Image Registry', 'dlisproddockerrepo.azurecr.io/dlis/llm_framework_vllm:<tag>'],
        ['First Build Time', '~30 minutes'],
        ['Incremental Build Time', '~8 minutes (Siwen\'s experience)'],
    ]
)
add_para('Checking Build Status:', bold=True)
add_bullet('View build progress and logs on the ADO Pipelines page')
p = doc.paragraphs[-1]
p.clear()
p.add_run('View build progress and logs on the ')
add_hyperlink(p, 'ADO Pipelines page', OAAS_LLM_TEMPLATE_PIPELINES)
add_bullet('After a successful build, the Pipeline log outputs the final image tag')
add_bullet('Use this image tag for the ModelPath configuration in subsequent Polaris Jobs')

add_heading('6.3 PR Build (Optional)', 2)
add_para('If you need to merge into the main branch (e.g., for general feature improvements), submitting a PR also triggers a build. After PR merge, the main branch builds an official version.')
add_note('For project-specific model code, building on a personal branch is usually sufficient \u2014 no need to merge into main.')

# ========== Section 7: Polaris Testing ==========
add_heading('7. Step 5: Polaris Testing', 1)
add_para('After image build and data migration are complete, use a Polaris Job for testing to verify model output and latency meet expectations before DLIS production deployment.')

add_heading('7.1 Polaris Job Configuration', 2)
add_table(
    ['Field', 'Example Value', 'Description'],
    [
        ['ModelPath', 'docker-repo://dlisproddockerrepo.azurecr.io/dlis/llm_framework_vllm:<tag>', 'Image address'],
        ['ModelDataPath', 'abfs://dlisstore@dlisstoregen2.dfs.core.windows.net/.../complete.txt', 'Gen2 path'],
        ['Environment Variables', 'DLIS_MODEL_DATA_TARGET_PATH=/Model;GPU_MEMORY_UTILIZATION=0.7', ''],
        ['WaitingModelReadyInMin', '30', 'Model load timeout'],
    ]
)

add_heading('7.2 Polaris Job Status', 2)
add_image('zhahao_image010.jpg', 'Fig: Polaris Job submission page (Source: Hao Zhang doc)')
add_image('zhahao_image012.png', 'Fig: Polaris Job configuration parameters (1) (Source: Hao Zhang doc)')
add_image('zhahao_image014.png', 'Fig: Polaris Job configuration parameters (2) (Source: Hao Zhang doc)')
add_bullet('Instance Loading: 100% + Success means deployment is successful (no need to wait for Instance Activate)')
add_bullet('Usually completes within 30 minutes; you can do other work after submission (Siwen\'s experience)')

add_heading('7.3 Testing Verification Points', 2)
add_para('During Polaris testing, verify the following:', bold=True)
add_bullet('Output correctness: send test requests and check if model responses match expectations (format, content quality)')
add_bullet('Latency: record end-to-end response time, confirm it meets business SLA requirements')
add_bullet('Resource usage: observe GPU memory usage and CPU utilization are within reasonable bounds')
add_bullet('Stability: send multiple requests to confirm the service does not crash or return abnormal results')
add_note('If testing reveals issues, go back to Step 1/4 to modify code or config, rebuild the image, and resubmit Polaris testing.')
add_image('zhahao_image016.jpg', 'Fig: Polaris Job latency and QPS statistics after completion (Source: Hao Zhang doc)')

add_heading('7.4 Test Request Example', 2)
add_code('''import requests

response = requests.post(
    "https://WestUS2.bing.prod.dlis.binginternal.com/route/PicassoAdsCreative.<ModelName>",
    cert=("private1.cer", "private1.key"),
    json={"prompt": "test input"},
    headers={"Content-Type": "application/json"},
    verify=False,
)
print(f"Status: {response.status_code}")
print(f"Latency: {response.elapsed.total_seconds():.2f}s")
print(f"Response: {response.text[:500]}")
''')

# ========== Section 8: DLIS Production Deployment ==========
add_heading('8. Step 6: DLIS Production Deployment', 1)
add_para('After Polaris testing passes, proceed with DLIS production deployment.')

add_heading('8.1 Hardware Allocation', 2)
add_para('Choose appropriate hardware based on model requirements (Image Model Service group experience):')
add_table(
    ['Model Type', 'Recommended Hardware', 'Notes'],
    [
        ['Relevance Model', 'A100 / A100 Train', 'Large model inference requires high GPU memory'],
        ['Diversity Model', 'T4 / MIG7', 'Smaller models can use lower-spec GPUs'],
    ]
)
add_bullet('A100 machines may be resource-constrained; consider A100 Train as an alternative')
add_bullet('Check that instances have sufficient CPU resources (not just GPU)')
add_para_with_link('View available machines and quotas: ', 'DLIS Portal -> Quota V2', ONE_INFERENCE_PORTAL, ', select Namespace and expand:')
add_image('zhahao_image022.png', 'Fig: DLIS Portal Quota V2 page (Source: Hao Zhang doc)')
add_image('zhahao_image024.png', 'Fig: Expand Namespace to view available machine list (Source: Hao Zhang doc)')
add_image('zhahao_image026.png', 'Fig: Machine quota details (Source: Hao Zhang doc)')
add_doc_link('Source doc: How_to_Build_Your_Own_DLIS_Model.docx (Hao Zhang)', ZHAHAO_DOC_URL)

add_heading('8.2 Create DLIS Service', 2)
add_para_with_link('Reference: ', 'How_to_Build_Your_Own_DLIS_Model.docx (Step 8)', ZHAHAO_DOC_URL, ' to create a DLIS Service.')
add_para('Steps:', bold=True)
add_bullet('1. Open DLModelV2 - One Inference Portal, click New Model')
p = doc.paragraphs[-1]
p.clear()
p.add_run('1. Open ')
add_hyperlink(p, 'DLModelV2 - One Inference Portal', ONE_INFERENCE_PORTAL)
p.add_run(', click New Model')
add_image('zhahao_image018.png', 'Fig: One Inference Portal - click New Model (Source: Hao Zhang doc)')
add_bullet('2. Paste Polaris Job Id, configure each page:')
add_image('zhahao_image020.png', 'Fig: Configure Key, Hardware, General, ACL pages (Source: Hao Zhang doc)')
add_bullet('Key page: update Environment and Namespace')
add_bullet('Hardware page: select deployment target machine')
add_bullet('General page: fill in DRI contact, set min/max instance count, set Model Priority to Test (use Production for prod)')
add_bullet('ACL page: add access control ACL')

add_heading('8.3 ACL Configuration', 2)
add_para('The DLIS Service ACL string contains multiple certificate thumbprints and AAD application IDs, controlling who can call the service:')
add_code('*:Certificate://Thumbprint/02AAAAA5AD...,*:AAD://appid/dda2a640-...,\n*:Certificate://Microsoft/dlis.si.advisoraggregator.trafficmanager.net,...')
add_note('Incorrect ACL configuration will cause callers to receive 403 Forbidden. If requests are rejected after deployment, check ACL configuration first.')

add_bullet('3. Click VALIDATION to validate, then click SUBMIT to submit')
add_bullet('Prod deployment requires a bypass process (manual operation by someone), which can be a bottleneck (Siwen\'s experience)')

# ========== Section 9: Post-Deployment Verification ==========
add_heading('9. Step 7: Post-Deployment Verification', 1)
add_para('After successful DLIS production deployment, send requests to verify the service is working correctly.')

add_heading('9.1 Endpoint Naming Conventions', 2)
add_bullet('Names should be descriptive and stable (e.g., PicassoAdsCreative.ZImage-V1)')
add_bullet('Avoid using personal names as endpoint names')
add_bullet('SI and Prod should not share the same endpoint (this is a pilot blocker)')

add_heading('9.2 Request URL Format', 2)
add_table(
    ['Incorrect Format', 'Correct Format'],
    [
        ['http://WestUS2BE.bing.prod.dlis.binginternal.com:86/route/...', 'https://WestUS2.bing.prod.dlis.binginternal.com/route/...'],
        ['With :8888 suffix', 'No port suffix needed'],
        ['/routebatch/', '/route/'],
    ]
)

add_heading('9.3 Test Request Example', 2)
add_code('''import requests

response = requests.post(
    "https://WestUS2.bing.prod.dlis.binginternal.com/route/PicassoAdsCreative.<ModelName>",
    cert=("private1.cer", "private1.key"),
    json={"prompt": "test input"},
    headers={"Content-Type": "application/json"},
    verify=False,
)
''')
add_para('Client certificates (.cer + .key files) are required for authentication; plain curl will not work.')
add_link_note('See "Section 11: SI/Prod Environments & Certificate Management" for certificate details')
add_link_note('See "Section 12: Kusto Log Viewing & Debugging" to verify Kusto logs are written correctly')

# ========== Section 10 ==========
add_heading('10. Step 8: Polaris Job Optimization (Optional)', 1)
add_note('This step is optional. If skipped, contact Fang Zhang to submit a bypass job.')
add_heading('10.1 Quantization', 2)
add_bullet('Automatic optimization: via Polaris Job built-in quantization workflow')
add_bullet('Offline quantization (recommended): use AutoGPTQ, AutoAWQ, or llm-compressor')
add_bullet('After offline quantization, upload the quantized model to Cosmos; vLLM auto-detects quantization format from config.json')

add_heading('10.2 Other Optimizations', 2)
add_table(
    ['Optimization', 'Description'],
    [
        ['Async API Call', 'Asynchronous inference calls for improved throughput'],
        ['Continuous Batching', 'Continuous batching for significantly improved throughput'],
    ]
)

# ================================================================
# PART 2: ENVIRONMENT & OPS
# ================================================================
doc.add_page_break()
p = doc.add_paragraph()
run = p.add_run('Part 2: Environment & Operations')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)

# ========== Section 11 ==========
add_heading('11. SI/Prod Environments & Certificate Management', 1)

add_heading('11.1 SI and Prod Environment Separation', 2)
add_bullet('SI environment is for testing and validation; Prod environment is for production services')
add_bullet('Sharing the same endpoint between SI and Prod is a pilot blocker (confirmed by Image Model Service group); production launch requires separation')
add_bullet('Prod deployment requires a bypass process (manual operation), which can be a time bottleneck')

add_heading('11.2 Certificate Types and Usage', 2)
add_table(
    ['Certificate Type', 'Format', 'Purpose', 'Location'],
    [
        ['Client authentication cert', '.cer + .key', 'Call DLIS endpoint', 'Local machine or server'],
        ['Kusto log certificate', '.pfx', 'EventHub writes to Kusto logs', 'Cosmos /Model/ directory'],
        ['SSL Keys', '.cert + .key', 'Server-side SSL', 'Server-specific path'],
    ]
)

add_para('Certificate file location reference (provided by Siwen):')
add_bullet('SSL keys: 10.224.120.197 /home/siwen/relevance/deploy')
add_bullet('New certificates on Cosmos: cosmos09.osdinfra.net/.../ImgLPRelevance6/')

add_heading('11.3 Certificate Expiration Management', 2)
add_bullet('Certificates expire at the end of April 2026; new certificates use .pfx format')
add_bullet('Expired certificates will cause authentication failures; update in advance')
add_bullet('Recommend setting expiration reminders; update 2 weeks before expiry')

# ========== Section 12 ==========
add_heading('12. Kusto Log Viewing & Debugging', 1)

add_heading('12.1 Kusto Log Environment Routing', 2)
add_para('Key finding (confirmed by Siwen): the certificate determines which environment\'s logs you can see.', bold=True)
add_bullet('Using SI certificate -> can only see SI environment logs')
add_bullet('Using Prod certificate -> can only see Prod environment logs')
add_bullet('If namespace is Prod but certificate is SI, logs will be written to the wrong database')
add_note('Previously discovered SI Kusto logs erroneously written to Prod DB (Image Model Service group); ensure configuration consistency.')

add_heading('12.2 Kusto Queries', 2)
add_bullet('SI environment: bingadsppe.AdInsightMT')
add_bullet('Prod environment: bingads.BingAdsTracing')
add_table_with_links(
    ['Environment', 'Link'],
    [
        ['PROD Kusto', ('https://bingads.kusto.windows.net/', KUSTO_PROD_URL)],
        ['SI Kusto', ('https://bingadsppe.kusto.windows.net/', KUSTO_SI_URL)],
        ['DLIS Jarvis Dashboard (Prod)', ('DLIS Model Metrics | Jarvis', JARVIS_PROD_URL)],
        ['DLIS Jarvis Dashboard (SI)', ('DLIS Model Metrics | Jarvis (SI)', JARVIS_SI_URL)],
    ]
)
add_doc_link('Auto Image Service DLIS Documentation (ChunChen) \u2014 includes Kusto log analysis examples', CHUNCHEN_DOC_URL)
add_code('''appsvc_info | union appsvc_warn | union appsvc_err
| where Timestamp > ago(30min)
| where ApplicationName == 'ImgLPRelevanceModel'
''')
add_para('The ApplicationName value comes from the application_name field in the config file (confirmed by Siwen).')

add_heading('12.3 Polaris Log', 2)
add_bullet('Polaris log shows service startup logs but not container print logs (confirmed by Siwen)')
add_bullet('For more detailed logs, use Kusto or Central Log')
p = doc.paragraphs[-1]
p.clear()
p.add_run('For more detailed logs, use Kusto or ')
add_hyperlink(p, 'Central Log', DLIS_WIKI_CENTRAL_LOG)
add_link_note('See "Appendix E" for Central Log queries')

add_heading('12.4 Local Kusto Log Testing', 2)
add_para('Key principle: Integrate Kusto log output during the local development phase. Do not wait until online deployment to start using Kusto logging — you should be able to see Kusto logs during local Docker testing. This lets you catch log format, auth config, and EventHub connection issues early.', bold=True)
add_para('Local testing only requires: ① the correct PFX certificate file; ② the matching EventHub namespace config; ③ network access to the EventHub endpoint. With these three prerequisites met, logs from your local Docker container will be sent to Kusto in real-time and can be queried directly in Kusto Explorer.')
add_para('Test script using AAD Bearer token (MSAL + PFX certificate) authentication:')
add_code('''import msal, requests
from cryptography.hazmat.primitives.serialization.pkcs12 import load_key_and_certificates

# 1. Load PFX certificate -> get private_key, thumbprint, public_certificate
# 2. Get AAD Bearer token
app = msal.ConfidentialClientApplication(
    client_id=CLIENT_ID,
    authority=f"https://login.microsoftonline.com/{TENANT_ID}",
    client_credential={"private_key": ..., "thumbprint": ..., "public_certificate": ...},
)
result = app.acquire_token_for_client(scopes=[DLIS_SCOPE])

# 3. Send request with tracking_data
payload = {
    "prompt": "test", "width": 1344, "height": 768,
    "tracking_data": {
        "requestid": f"test-{uuid.uuid4().hex[:8]}",
        "trackingid": f"test-{uuid.uuid4().hex[:8]}",
        "callername": "local_test",
    },
}
resp = requests.post(API_URL, json=payload,
                     headers={"Authorization": f"Bearer {token}"},
                     verify=False, timeout=120)
''')

add_heading('12.5 Common Kusto Log Issues & Debugging Tips', 2)
add_para('The following lessons are summarized from real debugging experiences during ZImage, Gemma4, and other model deployments.')

add_heading('Issue 1: EventHub Auth Errors Silently Swallowed', 3)
add_para('Symptom: Local test returns results normally, but Kusto shows zero logs.')
add_para('Root cause: kusto_log.py catches all EventHub send exceptions with a bare try/except pass, making certificate or namespace misconfigurations completely invisible.')
add_code('''# ❌ Wrong — silently swallows auth failures
try:
    client.send(event_data)
except Exception:
    pass  # logs lost, no indication

# ✅ Correct — fail fast, expose config issues
try:
    client.send(event_data)
except Exception as e:
    logger.error(f"EventHub send failed: {e}", exc_info=True)
    raise  # first failure should be immediately visible''')

add_heading('Issue 2: record.msg vs record.getMessage()', 3)
add_para('Symptom: Kusto log messages show raw template strings (e.g., "%s loaded in %d seconds") instead of formatted values.')
add_para('Root cause: kusto_log.py uses record.msg which is the unformatted template. Use record.getMessage() to get the fully formatted string.')
add_code('''# ❌ record.msg → "Model %s loaded in %d seconds"
# ✅ record.getMessage() → "Model gemma4 loaded in 42 seconds"''')

add_heading('Issue 3: Kusto Logs Lost on Process Crash', 3)
add_para('Symptom: OOM or CUDA errors during model loading crash the process, but Kusto shows no error logs.')
add_para('Root cause: KustoHandler uses ScheduledBatchSender for periodic batch sending. When the process crashes, the scheduler thread dies with it and all buffered logs are lost.')
add_code('''# Solution: manually flush in crash handler
import signal, atexit

def flush_kusto_on_exit():
    for handler in logging.root.handlers:
        if hasattr(handler, 'flush'):
            handler.flush()

atexit.register(flush_kusto_on_exit)
signal.signal(signal.SIGTERM, lambda *_: (flush_kusto_on_exit(), sys.exit(1)))''')

add_heading('Issue 4: SI/Prod Certificate-Namespace Mismatch', 3)
add_para('Symptom: Log sending shows no errors (if exceptions are caught), but Kusto queries return nothing.')
add_para('Root cause: Using an SI certificate with a Prod EventHub namespace (or vice versa). Auth may pass but messages are routed to the wrong environment.')
add_bullet('SI environment: namespace has "si" suffix, use SI certificate')
add_bullet('Prod environment: namespace has no suffix, use Prod certificate')
add_para('Recommendation: Put environment config in settings.json and switch via environment variables — avoid hardcoding.', bold=True)

add_heading('Issue 5: Logger Level Defaults to WARNING', 3)
add_para('Symptom: Code has logger.info() calls, but Kusto only shows WARNING and above.')
add_para('Root cause: Python logging child loggers inherit the root logger\'s level (WARNING) by default. Without explicit configuration, all INFO and DEBUG logs are filtered out.')
add_code('''# Must explicitly set logger level
logger = logging.getLogger("dlis_model")
logger.setLevel(logging.INFO)  # defaults to WARNING if not set''')

add_heading('Issue 6: Validate Kusto Logs During Local Testing', 3)
add_para('Important: Do not defer Kusto log validation to the online deployment stage. Integrate Kusto logging during local Docker testing to ensure logs are sent to EventHub correctly.')
add_para('Prerequisites for local Kusto log testing:')
add_bullet('Correct environment PFX certificate file (SI or Prod)')
add_bullet('settings.json configured with the correct EventHub namespace and topic')
add_bullet('Network access to EventHub endpoint (corporate network or VPN)')
add_para('Validation steps:')
add_bullet('After starting local Docker, send a test request and check container logs for EventHub send success/failure output')
add_bullet('Simultaneously query the corresponding Kusto environment table to confirm logs have arrived (typically 1-2 minute delay)')
add_bullet('If network restrictions prevent EventHub access, temporarily use a console handler to verify log format, but do full EventHub validation as soon as possible')

add_heading('Issue 7: EventHub Four Topics Explained', 3)
add_para('DLIS EventHub provides four topics. Send logs to the correct topic based on log type:')
add_bullet('appsvc_info — general information logs (model loading, request processing, etc.)')
add_bullet('appsvc_warn — warning logs (non-fatal errors, performance degradation, etc.)')
add_bullet('appsvc_err — error logs (exceptions, crash info, etc.)')
add_bullet('appsvc_perf — performance logs (inference latency, throughput metrics, etc.)')
add_para('Note: If you only send to appsvc_info, querying the errors table in Kusto will return nothing.', bold=True)

# ========== Section 13 ==========
add_heading('13. Common Issues & Solutions', 1)

problems = [
    ('Issue 1: Container OOM Killed (CPU memory, not GPU)',
     'OaasWrapper cannot find _opt directory -> falls back to BaseLLM -> loads all weights on CPU -> exceeds CPU memory limit',
     'Ensure _opt directory structure is correct (Approach A), or use vllm.LLM() directly (Approach B, recommended)'),
    ('Issue 2: CUDA_VISIBLE_DEVICES UUID Format',
     'DLIS sets CUDA_VISIBLE_DEVICES to GPU UUID format; vLLM internal int() conversion fails',
     'In run.sh: unset CUDA_VISIBLE_DEVICES + in model.py: convert UUID to integer indices'),
    ('Issue 3: /Model is a Read-Only Filesystem',
     'DLIS mounts Cosmos data to /Model in read-only mode',
     'Writable Mirror approach: create mirror directory in /tmp, use symlinks to point to read-only model files'),
    ('Issue 4: Unable to find exposed port 8888',
     'Usually not a port issue but a model loading crash (OOM, etc.); HTTP server never started',
     'Fix model loading issues first; also ensure Dockerfile has EXPOSE 8888'),
    ('Issue 5: Pipeline Build pip install Timeout',
     'CI agent cannot directly connect to pypi.org',
     'Add Tsinghua PyPI mirror (see Appendix B)'),
    ('Issue 6: OaasWrapper _create_runner() Silently Swallows Exceptions',
     '_create_runner() has try/except that catches all exceptions, prints, then returns None',
     'Change to logger.error(exc_info=True), or use Approach B directly'),
    ('Issue 7: opt_type.txt Newline Character',
     'echo "llm" writes with trailing newline',
     'Use printf "llm" instead of echo "llm"'),
    ('Issue 8: Cosmos Mount Only Mounts Top-Level Directory',
     'Placing certificates in subdirectories causes "not found" errors',
     'Place certificates in the Cosmos root directory, not in subdirectories'),
    ('Issue 9: Docker in Logs Mismatches Config',
     'Config file does not match the actual running image',
     'Check that ModelPath in Polaris Job config points to the correct image tag'),
    ('Issue 10: Local Test Returns Results but Kusto Logs Are Empty',
     'Using the wrong environment\'s certificate (e.g., namespace is Prod but certificate is SI)',
     'Ensure namespace and certificate environment are consistent (see Sections 11 and 12)'),
]

for title, cause, fix in problems:
    add_heading(title, 2)
    add_para(f'Root Cause: {cause}')
    add_para(f'Solution: {fix}')

# ================================================================
# APPENDICES
# ================================================================
doc.add_page_break()
p = doc.add_paragraph()
run = p.add_run('Appendices')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)

# ========== Appendix A ==========
add_heading('Appendix A: Base Docker Image Selection & Updates', 1)

add_heading('A.1 Fast Build Approach (Recommended): Dockerfile_vllm_fast', 2)
add_code('''FROM vllm/vllm-openai:latest
# Already includes vllm, torch, transformers - no compilation needed
# If a specific transformers version is needed:
RUN python3 -m pip install transformers==5.5.3
''')
add_bullet('Pros: Build < 1 second, includes pre-compiled vllm + torch')
add_bullet('Cons: latest version is uncontrolled, may be updated upstream')

add_heading('A.2 Full Build Approach: Dockerfile_vllm_0.10.0', 2)
add_code('''FROM nvidia/cuda:12.8.1-devel-ubuntu22.04

# Key: install the correct version of torch first, then install vllm
RUN pip install torch==2.10.0 torchvision==0.25.0 \\
    --index-url https://download.pytorch.org/whl/cu128
RUN pip install vllm==0.19.0
''')

add_heading('A.3 Version Compatibility Experience', 2)
add_table(
    ['Issue', 'Root Cause', 'Solution'],
    [
        ['vllm/_C.abi3.so: undefined symbol', 'torch ABI mismatch', 'Install torch first; do not add torch/torchvision to requirements'],
        ['torchvision::nms does not exist', 'Version mismatch', 'Ensure torchvision matches torch version'],
        ['num_scheduler_steps not accepted', 'Removed in vllm 0.19.0', 'Remove from vllm_runner.py'],
        ['Gemma4VideoProcessor requires Torchvision', 'torchvision uninstalled', 'Reinstall the correct version'],
    ]
)
add_para('Principle: Do not add torch, torchvision, or transformers to requirements-vllm.txt; let Dockerfile manage them.', bold=True)

add_heading('A.4 Local Manual Build Test', 2)
add_code('''cd OaaS_LLMTemplate
IMAGE_TAG="local-test"

# Block 1: Build base image
sudo docker build -t my-vllm-base:$IMAGE_TAG \\
    --file pipeline/Dockerfile_vllm_fast pipeline/

# Block 2: Install OaaS code and dependencies
sudo docker run -d --name temp my-vllm-base:$IMAGE_TAG sleep infinity
sudo docker cp . temp:/
sudo docker exec temp chmod +x /dlis_model/run.sh /dlis_model/async_run.sh /LLMModelOptimization.sh
sudo docker exec temp python3 -m pip install -r /requirements-common.txt
sudo docker exec temp python3 -m pip install -r /requirements-vllm.txt
sudo docker exec temp python3 -m pip install -e /
sudo docker commit temp my-final:$IMAGE_TAG
sudo docker rm -f temp
''')

# ========== Appendix B ==========
add_heading('Appendix B: Tsinghua Mirror Configuration', 1)
add_para('DLIS CI pipeline agents and Docker builds may not be able to directly connect to pypi.org.')
add_heading('B.1 CI Pipeline (azure-pipelines-unified.yml)', 2)
add_code('''variables:
  PIP_INDEX_URL: https://pypi.tuna.tsinghua.edu.cn/simple
  UV_INDEX_URL: https://pypi.tuna.tsinghua.edu.cn/simple
''')
add_heading('B.2 Docker Build (build_vllm_image.sh)', 2)
add_code('''docker build \\
  --build-arg PIP_INDEX_URL=https://pypi.tuna.tsinghua.edu.cn/simple \\
  --build-arg UV_INDEX_URL=https://pypi.tuna.tsinghua.edu.cn/simple \\
  --build-arg PIP_EXTRA_INDEX_URL=https://pypi.org/simple \\
  ...

# docker exec installation phase:
PIP_ARGS="-i https://pypi.tuna.tsinghua.edu.cn/simple --extra-index-url https://pypi.org/simple"
docker exec container pip install $PIP_ARGS -r requirements.txt
''')
add_heading('B.3 Dockerfile Parameter Reception', 2)
add_code('''ARG PIP_INDEX_URL
ARG UV_INDEX_URL
ARG PIP_EXTRA_INDEX_URL
ARG UV_EXTRA_INDEX_URL
''')
add_note('Note: Cannot use docker exec -e environment variable approach for mirror params; must use --build-arg or CLI arguments.')

# ========== Appendix C ==========
add_heading('Appendix C: Model Code Writing Guide', 1)

add_heading('C.1 model.py Core Structure', 2)
add_code('''class ModelImp:
    def __init__(self):
        # Model initialization: load engine, configure parameters
        pass

    def Eval(self, data):
        # Single inference
        pass

    def EvalBatch(self, data_list):
        # Batch inference
        pass
''')

add_heading('C.2 Approach A: Using OaasWrapper (For Simple LLM Inference)', 2)
add_code('''from llm_opt.oaas_wrapper_v2 import OaasWrapper

class ModelImp:
    def __init__(self):
        self.oaas_wrapper = OaasWrapper("model_dir_name", is_llm_model=True)

    def Eval(self, data):
        prompts = preprocess(data)
        outputs = self.oaas_wrapper.run(prompts)
        return postprocess(outputs)
''')
add_bullet('Requires opt_type.txt and best_setting.json in the _opt directory')
add_note('Note: Do not set the quantization field in best_setting.json. vLLM auto-detects from model config.json.')

add_heading('C.3 Approach B: Using vLLM Directly (Recommended, More Control)', 2)
add_code('''from vllm import LLM, SamplingParams

class ModelImp:
    def __init__(self):
        # CUDA_VISIBLE_DEVICES UUID fix
        cuda_env = os.environ.get("CUDA_VISIBLE_DEVICES", "")
        if cuda_env and not cuda_env.replace(",", "").isdigit():
            gpu_count = len(cuda_env.split(","))
            os.environ["CUDA_VISIBLE_DEVICES"] = ",".join(str(i) for i in range(gpu_count))

        self.llm = LLM(
            model="/Model/model-name",
            tensor_parallel_size=1,
            gpu_memory_utilization=float(os.environ.get("GPU_MEMORY_UTILIZATION", "0.9")),
            trust_remote_code=True, dtype="auto", max_model_len=8192,
            enable_prefix_caching=True,
        )

        self.sampling_params = SamplingParams(
            temperature=1.0, top_p=0.95, top_k=64,
            max_tokens=128, stop=["</end_token>"],
        )

    def Eval(self, data):
        prompts = preprocess(data)
        outputs = self.llm.generate(prompts, self.sampling_params)
        return [o.outputs[0].text for o in outputs]
''')
add_para('Advantages of Approach B:', bold=True)
add_bullet('Removes the OaasWrapper middle layer; initialization failures are reported directly')
add_bullet('No need for _opt directory or best_setting.json')
add_bullet('Inference parameters are written directly in code, transparent and controllable')
add_note('Siwen also uses vllm.LLM directly instead of OaasWrapper; this has become team consensus.')

add_heading('C.4 Non-LLM Models (e.g., ZImage Diffusion)', 2)
add_code('''class ModelImp:
    def __init__(self):
        from diffusers import ZImagePipeline
        model_path = os.environ.get("ZIMAGE_MODEL_PATH", "/Model/model_name")
        self.pipe = ZImagePipeline.from_pretrained(model_path, torch_dtype=torch.bfloat16)
        self.pipe.to("cuda:0")

    def Eval(self, data):
        result = self.pipe(prompt=data["prompt"], width=data["width"], height=data["height"])
        ...
''')

add_heading('C.5 Multi-Step Inference (e.g., Gemma4 Two-Step)', 2)
add_code('''def Eval(self, data):
    # Step 1: Generate scene concepts
    step1_outputs = self._run_vllm(step1_prompts, self.step1_params)

    # Step 2: Expand into detailed prompts
    step2_outputs = self._run_vllm(step2_prompts, self.step2_params)
    return self.processor.postprocess(step2_outputs)
''')
add_note('Key: Step 1 and Step 2 should use independent SamplingParams (different max_tokens and stop tokens).')

add_heading('C.6 Tokenizer Thinking Mode Issue', 2)
add_para('Some quantized model tokenizers have a thinking prefix built into chat_template that needs patching:')
add_code('''tokenizer = self.llm.get_tokenizer()
if hasattr(tokenizer, 'chat_template') and '<|channel>thought' in (tokenizer.chat_template or ''):
    tokenizer.chat_template = tokenizer.chat_template.replace(
        "<|channel>thought\\n<channel|>", ""
    )
''')

# ========== Appendix D ==========
add_heading('Appendix D: External settings.json Configuration Override', 1)

add_heading('D.1 Background', 2)
add_para('If configurations in config.py are hardcoded, switching between SI/Prod environments requires code changes and image rebuilds. Reference Hanbang\'s implementation (user/hanbangliang/img-outpainting-v1 branch) for overriding via an external JSON file on Cosmos.')

add_heading('D.2 Implementation', 2)
add_code('''from pydantic_settings import BaseSettings, SettingsConfigDict

DEFAULT_SETTINGS_JSON_PATH = os.environ.get("SETTINGS_JSON_PATH", "/Model/settings.json")

class Settings(BaseSettings):
    model_config = SettingsConfigDict(extra="ignore")
    eventhub_namespace: str = "aggregation-logging.servicebus.windows.net"
    certificate_path: str = os.path.join("/Model", "AggSvcAuthCert-prod.pfx")
    ...
''')

add_heading('D.3 Configuration Priority', 2)
add_bullet('1. init kwargs (explicit code parameters)')
add_bullet('2. Environment variables (e.g., EVENTHUB_NAMESPACE=...)')
add_bullet('3. JSON file (/Model/settings.json)')
add_bullet('4. Code default values')

add_heading('D.4 Usage', 2)
add_para('Place settings.json in the Cosmos root directory, writing only fields that need overriding:')
add_code('''{"eventhub_namespace": "aggregation-si-logging.servicebus.windows.net",
 "certificate_path": "/Model/AggSvcAuthCert-si.pfx"}
''')
add_para('Without settings.json, code defaults are used (Prod environment).')

add_heading('D.5 EventHub Credential Fault Tolerance', 2)
add_code('''def _try_get_credential(tenant_id):
    try:
        return CertificateCredential(...)
    except Exception as e:
        logger.info("EventHub credential unavailable, kusto logs local-only")
        return None
''')
add_para('Certificate loading failure should not block container startup; silently degrade to local logging.')

# ========== Appendix E ==========
add_heading('Appendix E: Central Log Debugging Tool', 1)
add_para('Use Central Log to debug Polaris Jobs and view container output logs (recommended by Desheng):')
add_code('''SELECT machine_name, log_level, log_time, description
FROM dlissensitivelog
WHERE file_name LIKE 'DLMSUserLog_ContainerOutput%.log'
  AND log_time BETWEEN TIMESTAMP '2026-02-27 00:00:00'
                   AND TIMESTAMP '2026-02-28 18:00:00'
  AND machine_name = '<your_machine_name>'
LIMIT 10000;
''')

# ========== Appendix F ==========
add_heading('Appendix F: Reference Links', 1)
add_table_with_links(
    ['Resource', 'Link / Description'],
    [
        ['DLIS LLM Engine Official Docs', ('Optimizing a LLM model with DLIS LLM engine', DLIS_WIKI_LLM_ENGINE)],
        ['OaaS Template Repo', ('OaaS_LLMTemplate (ADO Git)', OAAS_LLM_TEMPLATE_REPO)],
        ['Gen1->Gen2 Data Migration Tools', ('Data Transfer Tools (DLIS Wiki)', DLIS_WIKI_DATA_TRANSFER)],
        ['Central Log Query Guide', ('How to Use Central Log (DLIS Wiki)', DLIS_WIKI_CENTRAL_LOG)],
        ['DLIS Model Building Guide (Hao Zhang)', ('How_to_Build_Your_Own_DLIS_Model.docx', ZHAHAO_DOC_URL)],
        ['OaaS Deployment Guide v2 (ChangXu)', ('DLIS_Model_DeploymentWith_OaaS_v2.docx', XUCHA_DOC_URL)],
        ['Kusto Logs (SI)', ('bingadsppe.AdInsightMT | Azure Data Explorer', KUSTO_SI_URL)],
        ['Kusto Logs (Prod)', ('bingads.BingAdsTracing | Azure Data Explorer', KUSTO_PROD_URL)],
        ['Auto Image Service DLIS Doc (ChunChen)', ('Auto Image Service \u2014 DLIS Model Documentation', CHUNCHEN_DOC_URL)],
        ['DLIS Deployment Walkthrough (Siwen & Desheng)', ('Call with Desheng Cui \u2014 2026-03-26 Recording', SIWEN_DESHENG_RECORDING)],
    ]
)
add_doc_link('DLIS_Model_DeploymentWith_OaaS_v2.docx (ChangXu)', XUCHA_DOC_URL)
add_doc_link('How_to_Build_Your_Own_DLIS_Model.docx (Hao Zhang)', ZHAHAO_DOC_URL)
add_doc_link('Auto Image Service \u2014 DLIS Model Documentation (ChunChen)', CHUNCHEN_DOC_URL)
add_doc_link('Call with Desheng Cui \u2014 2026-03-26 DLIS Deployment Walkthrough (Siwen & Desheng)', SIWEN_DESHENG_RECORDING)

# ========== Appendix G ==========
add_heading('Appendix G: Gemma4 DLIS Deployment Iteration Summary', 1)
add_table(
    ['Deployment #', 'Issue', 'Fix'],
    [
        ['#1-#3', '_opt directory invisible -> OOM', 'Confirmed Cosmos sync cannot resolve'],
        ['#4', '/Model read-only, cannot create _opt', 'Writable mirror approach'],
        ['#5', 'vLLM init failure silently swallowed', '_create_runner() changed to raise exception'],
        ['#6', 'Changed to raise -> crash loop', 'Reverted to return None + logger.error'],
        ['#7', 'CUDA_VISIBLE_DEVICES UUID format', 'UUID -> integer index conversion'],
        ['#8', 'Cosmos directory + stale model.py', 'New flat directory + removed OaasWrapper'],
    ]
)
add_para('Core lesson: OaasWrapper middle layer complexity far outweighs its benefits. Recommend using vllm.LLM() directly (Approach B).', bold=True)

# ========== Appendix H ==========
add_heading('Appendix H: Revision History', 1)
add_table(
    ['Version', 'Date', 'Author', 'Description'],
    [
        ['v2', '2026-02-26', 'ChangXu', 'Initial document, complete deployment guide'],
        ['v4', '2026-04-22', 'Jinjin Chen', 'Consolidated Gemma4/ZImage hands-on experience'],
        ['v5', '2026-04-22', 'Jinjin Chen', 'Consolidated ChangXu v2, settings.json, Kusto testing'],
        ['v5.1', '2026-04-22', 'Jinjin Chen', 'Restructured: main flow + appendices; consolidated team Teams chat insights'],
        ['v5.2', '2026-04-23', 'Jinjin Chen', 'English version; added Kusto/Jarvis links and ChunChen doc reference'],
    ]
)

# Save
output_path = r"C:\Users\jinjinchen\OneDrive - Microsoft\DLIS_Model_Deployment_Guide_v5_5_EN.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
