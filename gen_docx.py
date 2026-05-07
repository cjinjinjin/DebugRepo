#!/usr/bin/env python3
"""Generate DLIS_Model_Deployment_Guide_v5.docx — restructured with main flow + appendices."""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0, 51, 102)

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    return p

def add_link_note(text):
    """Add a quick-link reference note (bold + colored)"""
    p = doc.add_paragraph()
    run = p.add_run('→ ' + text)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.font.size = Pt(10)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    shd = run._element.makeelement(qn('w:shd'), {})
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F0F0F0')
    rpr = run._element.find(qn('w:rPr'))
    if rpr is None:
        rpr = run._element.makeelement(qn('w:rPr'), {})
        run._element.insert(0, rpr)
    rpr.append(shd)
    return p

def add_table_with_links(headers, rows):
    """Add a table where cell values can be (text, url) tuples for hyperlinks."""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            if isinstance(val, tuple):
                text, url = val
                cell.text = ''
                p = cell.paragraphs[0]
                add_hyperlink(p, text, url)
            else:
                cell.text = str(val)
    doc.add_paragraph()

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri+1].cells[ci].text = str(val)
    doc.add_paragraph()

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

# -- Image paths --
IMG_DIR = r'C:\Users\jinjinchen\OneDrive - Microsoft\doc_images'

def add_image(filename, caption=None, width=Inches(5.5)):
    """Add an image with optional caption."""
    import os
    img_path = os.path.join(IMG_DIR, filename)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=width)
        if caption:
            p = doc.add_paragraph()
            run = p.add_run(caption)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(120, 120, 120)
    else:
        add_note(f'[图片缺失: {filename}]')

# -- Document links --
XUCHA_DOC_URL = 'https://microsoftapc-my.sharepoint.com/:w:/g/personal/xucha_microsoft_com/IQCgQ8ufJRUWSoedwJNRll1yAc-aqBPPV5Rqmg4JjANvS2E'
ZHAHAO_DOC_URL = 'https://microsoftapc-my.sharepoint.com/:w:/g/personal/zhahao_microsoft_com/IQDV9SZPUx2jRbIULT1vcKGTAXmrH1bD-lcaqsz2L80tl_0'
CHUNCHEN_DOC_URL = 'https://microsoftapc-my.sharepoint.com/:w:/g/personal/chunchen_microsoft_com/IQDG4fXi2Zd3RqUUqu9dTNLqAaeNMD_4epy-2DcbxOt1pq8'

# -- External resource URLs --
DLIS_WIKI_LLM_ENGINE = 'https://dlisinfrawiki.azurewebsites.net/wiki/Contents/DLISLLMEngine/Optimizing+a+LLM+model+with+DLIS+LLM+engine.html'
DLIS_WIKI_DATA_TRANSFER = 'https://dlisinfrawiki.azurewebsites.net/wiki/Contents/ModelRepository/Data+Transfer+tools.html'
DLIS_WIKI_CENTRAL_LOG = 'https://dlisinfrawiki.azurewebsites.net/wiki/Contents/Debugging/How+to+use+Central+Log.html'
ONE_INFERENCE_PORTAL = 'https://dlis-portal.microsoft.com/'
OAAS_LLM_TEMPLATE_REPO = 'https://msasg.visualstudio.com/Bing_and_IPG/_git/OaaS_LLMTemplate'
OAAS_LLM_TEMPLATE_PIPELINES = 'https://msasg.visualstudio.com/Bing_and_IPG/_build?definitionScope=%5COaaS_LLMTemplate'
DLIS_COPY_PIPELINE = 'https://msasg.visualstudio.com/Bing_and_IPG/_build?definitionId=43927'
COSMOS_REPO_URL = 'https://cosmos09.osdinfra.net:443/cosmos/DLISModelRepository/'
KUSTO_SI_URL = 'https://bingadsppe.kusto.windows.net/'
KUSTO_PROD_URL = 'https://bingads.kusto.windows.net/'
JARVIS_PROD_URL = 'https://jarvis-west.dc.ad.msft.net/dashboard/DLIS-Model-Metrics'
JARVIS_SI_URL = 'https://jarvis-west.dc.ad.msft.net/dashboard/DLIS-Model-Metrics-SI'
SIWEN_DESHENG_RECORDING = 'https://teams.microsoft.com/l/meetingrecap?driveId=b%21FSzYuSuEp02FCYr9pdB1t_R60kGlONJJht32xiFgk6Df_y20EhFwRaIwLXCzCWLd&driveItemId=015S6O6HUNTJFSKRTY3NH24GGIFES2P7NK&sitePath=https%3A%2F%2Fmicrosoftapc-my.sharepoint.com%2Fpersonal%2Fsiwenzhu_microsoft_com%2FDocuments%2FRecordings%2FCall+with+Desheng+Cui-20260326_203633-Meeting+Recording.mp4%3Fweb%3D1&fileUrl=https%3A%2F%2Fmicrosoftapc-my.sharepoint.com%2Fpersonal%2Fsiwenzhu_microsoft_com%2FDocuments%2FRecordings%2FCall+with+Desheng+Cui-20260326_203633-Meeting+Recording.mp4%3Fweb%3D1&threadId=19%3A01d7b57a-2d8a-45ef-b6eb-9d9cd1c21105_0b99169c-642e-4bbd-afec-02523f1445b5%40unq.gbl.spaces&callId=a347d990-0384-4b69-a3a9-00885914510e&threadType=OneOnOneChat&meetingType=Unknown&subType=RecapSharingLink_RecapCore'

def add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')  # 10.5pt
    rPr.append(sz)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph

def add_para_with_link(prefix, link_text, url, suffix=''):
    """Add a paragraph with inline hyperlink."""
    p = doc.add_paragraph()
    if prefix:
        p.add_run(prefix)
    add_hyperlink(p, link_text, url)
    if suffix:
        p.add_run(suffix)
    return p

def add_note_with_link(prefix, link_text, url, suffix=''):
    """Add a note-styled paragraph with inline hyperlink."""
    p = doc.add_paragraph()
    if prefix:
        run = p.add_run(prefix)
        run.italic = True
        run.font.color.rgb = RGBColor(0, 102, 153)
    add_hyperlink(p, link_text, url)
    if suffix:
        run = p.add_run(suffix)
        run.italic = True
        run.font.color.rgb = RGBColor(0, 102, 153)
    return p

def add_doc_link(label, url):
    """Add a document reference link with clickable hyperlink."""
    p = doc.add_paragraph()
    run = p.add_run('📎 ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.font.size = Pt(10)
    add_hyperlink(p, label, url)

# ================================================================
# TITLE
# ================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DLIS 模型部署指南（基于 OaaS_LLMTemplate）')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('版本 v5.1 — 2026-04-22\n整合 Gemma4、ZImage 部署实战经验、ChangXu v2 文档及团队 Teams 聊天经验')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ================================================================
# TOC
# ================================================================
add_heading('目录', 1)
toc_items = [
    '第一部分：部署主流程',
    '  1. 简介',
    '  2. 部署流程概览',
    '  3. Step 1：本地开发与测试',
    '  4. Step 2：上传 Checkpoint 到 Gen1',
    '  5. Step 3：Gen1 → Gen2 数据迁移',
    '  6. Step 4：PR 提交与 CI 自动构建镜像',
    '  7. Step 5：Polaris 测试',
    '  8. Step 6：DLIS 正式部署',
    '  9. Step 7：部署后验证与测试',
    '  10. Step 8：Polaris Job 优化（可选）',
    '',
    '第二部分：环境与运维',
    '  11. SI/Prod 环境与证书管理',
    '  12. Kusto 日志查看与调试',
    '  13. 常见问题与解决方案',
    '',
    '附录',
    '  A. Base Docker 镜像选择与更新',
    '  B. 清华镜像源配置',
    '  C. Model 代码编写指南',
    '  D. 外部 settings.json 配置覆盖',
    '  E. Central Log 调试工具',
    '  F. 参考链接',
    '  G. Gemma4 部署迭代总结',
    '  H. 修订历史',
]
for item in toc_items:
    if item == '':
        doc.add_paragraph()
    elif item.startswith('第') or item.startswith('附录'):
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.bold = True
    else:
        doc.add_paragraph(item.strip(), style='List Bullet')

doc.add_page_break()

# ================================================================
# PART 1: MAIN FLOW
# ================================================================
p = doc.add_paragraph()
run = p.add_run('第一部分：部署主流程')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)

# ========== Section 1 ==========
add_heading('1. 简介', 1)
add_heading('1.1 本文档用途', 2)
add_para('本文档为部署 LLM / 多模态 / Diffusion 等模型到 DLIS（Deep Learning Inference Service）平台提供完整指导。')

add_heading('1.2 什么是 OaaS？', 2)
add_para('OaaS（Optimization as a Service）是 DLIS 上部署 LLM 的综合框架，提供：')
add_bullet('自动化模型优化：量化、编译、优化的标准化工作流')
add_bullet('多后端支持：vLLM 和 TensorRT-LLM 推理引擎')
add_bullet('生产就绪模板：预配置模板，快速部署')
add_bullet('批处理支持：高吞吐量场景下的高效批量推理')
add_note_with_link('📘 官方文档：', 'Optimizing a LLM model with DLIS LLM engine（DLIS Wiki）', DLIS_WIKI_LLM_ENGINE)

# ========== Section 2 ==========
add_heading('2. 部署流程概览', 1)
add_code('''  ┌──────────────────────────────────┐
  │ Step 1: 本地开发 & Docker 测试    │  ← 先验证功能正确性
  └───────────────┬────────────────┘
                  │
     ┌────────────┴────────────┐
     ▼                         ▼          （可同步进行）
  ┌────────────────────┐  ┌─────────────┐
  │Step 2: 上传 ckpt   │  │Step 4:      │
  │     到 Gen1        │  │PR & CI      │
  │        ↓           │  │构建镜像     │
  │Step 3: Gen1→Gen2   │  │             │
  │     数据迁移       │  │             │
  └─────────┬──────────┘  └──────┬──────┘
            └────────────┬───────┘
                         ▼
  ┌──────────────────────────────────┐
  │ Step 5: Polaris 测试              │  ← 检查输出和延迟；DLIS 从 Gen2 读取模型数据
  └───────────────┬────────────────┘
                  ▼
  ┌──────────────────────────────────┐
  │ Step 6: DLIS 正式部署             │
  └───────────────┬────────────────┘
                  ▼
  ┌──────────────────────────────────┐
  │ Step 7: 验证 & 监控               │
  └──────────────────────────────────┘''')

add_para('关键原则：', bold=True)
add_bullet('先本地验证，再上线。本地 Docker 测试通过后再上传数据和构建镜像')
add_bullet('本地测试阶段就应集成 Kusto 日志输出。不要等到线上部署才接入 Kusto——本地 Docker 测试时配好证书和 EventHub，即可实时在 Kusto 中查看日志，提前发现日志格式、认证、连接等问题')
add_bullet('DLIS 读取的模型数据来自 Gen2（dlisstoregen2.dfs.core.windows.net），不是 Gen1')
add_bullet('镜像通过 OaaS_LLMTemplate 仓库的 CI pipeline 自动构建')
p = doc.paragraphs[-1]
p.clear()
p.add_run('镜像通过 ')
add_hyperlink(p, 'OaaS_LLMTemplate 仓库', OAAS_LLM_TEMPLATE_REPO)
p.add_run(' 的 CI pipeline 自动构建')
add_bullet('Step 2+3（上传 ckpt + 数据迁移）和 Step 4（镜像构建）可以同步进行，互不依赖')

# ========== Section 3 ==========
add_heading('3. Step 1：本地开发与测试', 1)
add_heading('3.1 开发流程概览', 2)
add_para_with_link('在 ', 'OaaS_LLMTemplate 仓库', OAAS_LLM_TEMPLATE_REPO, ' 创建个人分支（如 jinjinchen/ZImage-v1），完成以下工作后再推送代码触发 CI 构建：')
add_bullet('1. 修改模型代码（model.py、dlis_inter.py 等）')
add_bullet('2. 评估是否需要定制 OaaS 模板（多模态支持、自定义格式等）')
add_bullet('3. 选择合适的 Dockerfile（快速迭代 vs 完整构建）')
add_bullet('4. 本地 Docker 构建并测试，确认功能正确')

add_heading('3.2 代码文件说明', 2)
add_para_with_link('在 ', 'OaaS_LLMTemplate 仓库', OAAS_LLM_TEMPLATE_REPO, ' 创建个人分支（如 jinjinchen/ZImage-v1），修改以下文件：')
add_table(
    ['文件', '说明'],
    [
        ['dlis_model/model/model.py', '模型初始化 + 推理逻辑（核心文件）'],
        ['dlis_model/model/dlis_inter.py', '预处理/后处理，实现 PreAndPostProcessor 类'],
        ['dlis_model/http_server.py', 'HTTP 服务（如需自定义格式）'],
        ['requirements-vllm.txt', 'Python 依赖'],
    ]
)
add_link_note('model.py 编写详见「附录 C：Model 代码编写指南」')

add_heading('3.3 OaaS 模板定制（可选）', 2)
add_note_with_link('模板定制是可选步骤。先检查 ', 'OaaS LLM Template 仓库', OAAS_LLM_TEMPLATE_REPO, ' 原始模板是否已满足需求。对于纯文本 LLM 推理任务，通常无需定制。')
add_table(
    ['定制需求', '说明', '修改文件'],
    [
        ['多模态 vLLM 支持', '原始模板仅支持文本输入', 'model.py'],
        ['图片传输格式', '改为 multipart/form-data 更高效', 'http_server.py'],
        ['额外依赖包', '如 diffusers、Pillow 等', 'requirements-vllm.txt'],
        ['非 LLM 模型', '如 Diffusion 模型（ZImage）', 'model.py'],
    ]
)

add_heading('3.4 两种 Dockerfile 选择', 2)
add_table(
    ['Dockerfile', '基础镜像', '构建时间', '适用场景'],
    [
        ['Dockerfile_vllm_0.10.0', 'nvidia/cuda:12.8.1-devel-ubuntu22.04', '几十分钟', '需要特定 vLLM/torch 版本'],
        ['Dockerfile_vllm_fast', 'vllm/vllm-openai:latest', '< 1 秒', '快速迭代'],
    ]
)
add_link_note('详细版本兼容性和构建问题见「附录 A：Base Docker 镜像选择与更新」')

add_para('Docker Build 常见网络问题：', bold=True)
add_bullet('apt-get install 失败：CI build agent 无法连接 archive.ubuntu.com')
add_bullet('pip install 超时：pypi.org 不可达')
add_link_note('解决方案见「附录 B：清华镜像源配置」')

add_heading('3.5 本地 Docker 测试', 2)
add_code('''# 1. 构建镜像
cd /path/to/OaaS_LLMTemplate
export SOURCE_BRANCH="test"
sudo bash pipeline/build_vllm_image.sh

# 2. 启动容器
IMAGE_TAG="<build_tag>"
sudo docker run -d --name model-test \\
  --gpus all \\
  -v /path/to/model_weights:/Model/model_name \\
  -p <host_port>:8888 \\
  <image_name>:$IMAGE_TAG \\
  /dlis_model/run.sh http

# 3. 测试请求
curl -X POST http://localhost:<host_port> \\
  -H "Content-Type: application/json" \\
  -d '{"prompt": "test input"}'
''')

add_para('注意事项：', bold=True)
add_bullet('如果某个 GPU 被占用，用 --gpus \'"device=N"\' 指定可用 GPU')
add_bullet('必须加 -p <host_port>:8888 端口映射，否则宿主机 curl 会 403')
add_bullet('vllm/vllm-openai 基础镜像的 entrypoint 是 vllm serve，如需交互式 bash 必须 --entrypoint bash')
add_bullet('如果 volume mount 了 Python 文件，注意 __pycache__ 可能导致旧代码被加载')

add_para('验证 Kusto 日志（建议在本地测试阶段完成）：', bold=True)
add_para('本地 Docker 容器启动后，除了验证推理功能，还应同步验证 Kusto 日志是否正常发送。这样可以在上线前提前发现证书、EventHub 配置等问题。')
add_bullet('1. 确保 Cosmos 目录中有正确环境的 PFX 证书（如 AggSvcAuthCert-si.pfx），并通过 volume mount 映射到容器内 /Model 目录')
add_bullet('2. 确保 settings.json 中配置了正确的 EventHub namespace 和 kusto_log 相关参数（证书路径、topic 等）')
add_bullet('3. 发送测试请求后，观察容器日志（docker logs）中是否有 EventHub 发送成功/失败的输出')
add_bullet('4. 在 Kusto Explorer 中查询对应环境的表，确认日志已到达（通常延迟 1-2 分钟）')
add_code('''# SI 环境 Kusto 查询示例
// Kusto cluster: https://bingadsppe.kusto.windows.net/
// Database: appsvc
appsvc_info
| where TIMESTAMP > ago(10m)
| where ModelName == "<your_model_name>"
| order by TIMESTAMP desc
| take 20''')
add_note('如果查询无结果，检查：① 证书环境是否与 namespace 匹配（SI 证书 + SI namespace）；② logger 级别是否设为 INFO；③ EventHub 发送是否有报错被静默吞掉。详见第 12 节。')

add_heading('3.6 离线测试（不启动 HTTP server）', 2)
add_code('''sudo docker run --rm -it --gpus all \\
  -v /path/to/model:/Model/model_name \\
  <image>:<tag> \\
  bash -c 'cd /dlis_model && ./run.sh offline /tmp/input.json /tmp/output.json'
''')

# ========== Section 4 ==========
add_heading('4. Step 2：上传 Checkpoint 到 Gen1', 1)
add_para('模型权重文件上传到 Gen1 Cosmos 存储。')
add_heading('4.1 上传地址', 2)
add_code('上传目标：https://cosmos09.osdinfra.net:443/cosmos/DLISModelRepository/local/<your-alias>/')
add_para('使用 Visual Studio Scope Extension 进行认证和上传。')
add_image('xucha_gen1_upload.png', '图：Gen1 Cosmos 上传目录示例（来源：ChangXu 文档）')
add_doc_link('原始文档：DLIS_Model_DeploymentWith_OaaS_v2.docx（ChangXu）', XUCHA_DOC_URL)

add_heading('4.2 目录结构（扁平化，推荐）', 2)
add_code('''<model-dir>/
├── model_name/               ← 模型权重文件夹
│   ├── config.json
│   ├── model-00001-of-N.safetensors
│   ├── tokenizer.json
│   └── ...
├── dlis_inter.py             ← 直接放根目录
├── settings.json             ← 配置覆盖（可选）
├── AggSvcAuthCert-prod.pfx   ← Kusto 证书（可选）
└── AggSvcAuthCert-si.pfx
''')

add_para('关键注意：', bold=True)
add_bullet('dlis_inter.py 必须放在 cosmos 根目录下，代码中 sys.path.append(\'/Model\') 只能找到根目录的文件')
add_bullet('不要在 cosmos 上放 model.py，用镜像自带的版本。cosmos 上的旧 model.py 会覆盖镜像里的新版本')
add_bullet('不要上传不需要的大文件（如 .tar 包），浪费同步时间')
add_bullet('证书等文件放 cosmos 根目录，不要嵌套在子文件夹中（DLIS 只挂载一级目录）')
add_link_note('settings.json 配置覆盖详见「附录 D：外部 settings.json 配置覆盖」')

# ========== Section 5 ==========
add_heading('5. Step 3：Gen1 → Gen2 数据迁移', 1)
add_note('注意：Step 2+3（上传 ckpt + 数据迁移）和 Step 4（镜像构建）可以同步进行，互不依赖。')
add_para('DLIS 部署时从 Gen2 读取模型数据，因此 Gen1 的数据需要迁移到 Gen2。')
add_para_with_link('参考 Wiki：', 'Data Transfer Tools（DLIS Wiki）', DLIS_WIKI_DATA_TRANSFER)
add_para_with_link('Gen1 → Gen2 迁移步骤（详见 ', 'How_to_Build_Your_Own_DLIS_Model.docx Step 6.2', ZHAHAO_DOC_URL, '）：')
add_bullet('1. 在 Repos 中创建分支')
add_bullet('2. 打开 DLIS copy pipeline，选择 View/Edit')
p = doc.paragraphs[-1]
p.clear()
p.add_run('2. 打开 ')
add_hyperlink(p, 'DLIS copy pipeline', DLIS_COPY_PIPELINE)
p.add_run('，选择 View/Edit')
add_image('zhahao_image002.png', '图：创建分支并打开 pipeline（来源：Hao Zhang 文档）')
add_bullet('3. 选择刚创建的分支，更新参数')
add_image('zhahao_image004.png', '图：选择分支并设置 pipeline 变量（来源：Hao Zhang 文档）')
add_image('zhahao_image006.png', '图：配置迁移路径参数（来源：Hao Zhang 文档）')
add_bullet('4. 点击 Validate and Save 保存参数')
add_image('zhahao_image008.png', '图：保存并运行 pipeline（来源：Hao Zhang 文档）')
add_doc_link('原始文档：How_to_Build_Your_Own_DLIS_Model.docx（Hao Zhang）', ZHAHAO_DOC_URL)
add_code('Gen2 路径格式：\nabfs://dlisstore@dlisstoregen2.dfs.core.windows.net/dlismodelrepository-c09/local/users/<username>/<model-dir>/')

add_para('Gen2 验证（重要）：', bold=True)
add_bullet('迁移完成后，使用 SAW（Secure Admin Workstation）验证 Gen2 上文件是否完整（Desheng 经验）')
add_bullet('Gen2 上有文件不代表没有错误，需要确认文件大小和完整性')
add_bullet('ADL 数据迁移工具有坑，建议迁移后仔细检查')

add_para('ModelDataPath 机制：', bold=True)
add_bullet('ModelDataPath 指向 Gen2 上的某个文件（如 complete.txt），DLIS 实际会把该文件所在的整个父目录挂载到 /Model')
add_bullet('/Model 以只读方式挂载，无法在其中创建文件')
add_bullet('如果需要在运行时创建配置文件，必须用 writable mirror 方案（见常见问题）')

# ========== Section 6 ==========
add_heading('6. Step 4：PR 提交与 CI 自动构建镜像', 1)
add_note('注意：Step 2+3（上传 ckpt + 数据迁移）和 Step 4（镜像构建）可以同步进行，互不依赖。')
add_para_with_link('本地开发和测试完成后，将代码推送到 ', 'OaaS_LLMTemplate 仓库', OAAS_LLM_TEMPLATE_REPO, ' 的个人分支，CI pipeline 会自动构建 Docker 镜像。')

add_heading('6.1 创建分支并提交代码', 2)
add_code('''# 在 OaaS_LLMTemplate 仓库创建个人分支
git checkout -b <your-alias>/<model-name>   # 如 jinjinchen/ZImage-v1

# 提交本地修改
git add -A
git commit -m "Add <model-name> model support"
git push origin <your-alias>/<model-name>
''')

add_heading('6.2 CI Pipeline 自动构建', 2)
add_para('推送到分支后，CI pipeline 会自动触发镜像构建。无需手动操作。')
add_table(
    ['项目', '说明'],
    [
        ['触发方式', 'Push 到任意分支自动触发（包括非 main 分支）'],
        ['镜像 Tag 格式', 'YYYYMMDD-HHMM-<branch_name>（非 main 分支）'],
        ['镜像仓库', 'dlisproddockerrepo.azurecr.io/dlis/llm_framework_vllm:<tag>'],
        ['首次构建时间', '约 30 分钟'],
        ['增量构建时间', '约 8 分钟（Siwen 经验）'],
    ]
)
add_para('查看构建状态：', bold=True)
add_bullet('在 ADO Pipelines 页面查看构建进度和日志')
p = doc.paragraphs[-1]
p.clear()
p.add_run('在 ')
add_hyperlink(p, 'ADO Pipelines 页面', OAAS_LLM_TEMPLATE_PIPELINES)
p.add_run(' 查看构建进度和日志')
add_bullet('构建成功后，Pipeline 日志中会输出最终的镜像 tag')
add_bullet('将该镜像 tag 用于后续 Polaris Job 的 ModelPath 配置')

add_heading('6.3 PR 构建（可选）', 2)
add_para('如果需要合入 main 分支（如通用功能改进），提交 PR 也会触发构建。PR 合并后 main 分支会构建正式版本。')
add_note('对于项目专用的模型代码，通常只需在个人分支构建即可，无需合入 main。')


# ========== Section 7: Polaris 测试 ==========
add_heading('7. Step 5：Polaris 测试', 1)
add_para('镜像构建和数据迁移完成后，先通过 Polaris Job 进行测试，验证模型输出结果和延迟是否符合预期，再进行 DLIS 正式部署。')

add_heading('7.1 Polaris Job 配置', 2)
add_table(
    ['字段', '示例值', '说明'],
    [
        ['ModelPath', 'docker-repo://dlisproddockerrepo.azurecr.io/dlis/llm_framework_vllm:<tag>', '镜像地址'],
        ['ModelDataPath', 'abfs://dlisstore@dlisstoregen2.dfs.core.windows.net/.../complete.txt', 'Gen2 路径'],
        ['环境变量', 'DLIS_MODEL_DATA_TARGET_PATH=/Model;GPU_MEMORY_UTILIZATION=0.7', ''],
        ['WaitingModelReadyInMin', '30', '模型加载超时时间'],
    ]
)

add_heading('7.2 Polaris Job 状态说明', 2)
add_image('zhahao_image010.jpg', '图：Polaris Job 提交页面（来源：Hao Zhang 文档）')
add_image('zhahao_image012.png', '图：Polaris Job 配置参数（1）（来源：Hao Zhang 文档）')
add_image('zhahao_image014.png', '图：Polaris Job 配置参数（2）（来源：Hao Zhang 文档）')
add_bullet('Instance Loading: 100% + Success 就是部署成功（不需要等 Instance Activate）')
add_bullet('一般半小时以内完成，提交后可以去做别的事（Siwen 经验）')

add_heading('7.3 测试验证要点', 2)
add_para('Polaris 测试阶段需要验证以下内容：', bold=True)
add_bullet('输出正确性：发送测试请求，检查模型返回结果是否符合预期（格式、内容质量）')
add_bullet('延迟（Latency）：记录端到端响应时间，确认是否满足业务 SLA 要求')
add_bullet('资源使用：观察 GPU 显存占用和 CPU 使用率是否在合理范围')
add_bullet('稳定性：发送多次请求，确认服务不会崩溃或返回异常结果')
add_note('如果测试发现问题，需要回到 Step 1/4 修改代码或配置后重新构建镜像，再次提交 Polaris 测试。')
add_image('zhahao_image016.jpg', '图：Polaris Job 完成后的 Latency 和 QPS 统计（来源：Hao Zhang 文档）')

add_heading('7.4 测试请求示例', 2)
add_code('''import requests

response = requests.post(
    "https://WestUS2.bing.prod.dlis.binginternal.com/route/PicassoAdsCreative.<ModelName>",
    cert=("private1.cer", "private1.key"),
    json={"prompt": "test input"},
    headers={"Content-Type": "application/json"},
    verify=False,
)
print(f"Status: {response.status_code}")
print(f"Latency: {response.elapsed.total_seconds():.2f}s")
print(f"Response: {response.text[:500]}")
''')

# ========== Section 8: DLIS 正式部署 ==========
add_heading('8. Step 6：DLIS 正式部署', 1)
add_para('Polaris 测试通过后，进行 DLIS 正式部署。')

add_heading('8.1 硬件分配', 2)
add_para('根据模型需求选择合适的硬件（Image Model Service 群组经验）：')
add_table(
    ['模型类型', '推荐硬件', '备注'],
    [
        ['Relevance Model', 'A100 / A100 Train', '大模型推理需要较大显存'],
        ['Diversity Model', 'T4 / MIG7', '较小模型可以使用低规格 GPU'],
    ]
)
add_bullet('A100 机器可能资源紧张，可以考虑 A100 Train 作为替代')
add_bullet('注意检查实例是否有足够的 CPU 资源（不只看 GPU）')
add_para_with_link('查看可用机器和配额：', 'DLIS Portal → Quota V2', ONE_INFERENCE_PORTAL, '，选择 Namespace 后展开查看：')
add_image('zhahao_image022.png', '图：DLIS Portal Quota V2 页面（来源：Hao Zhang 文档）')
add_image('zhahao_image024.png', '图：展开 Namespace 查看可用机器列表（来源：Hao Zhang 文档）')
add_image('zhahao_image026.png', '图：机器配额详情（来源：Hao Zhang 文档）')
add_doc_link('原始文档：How_to_Build_Your_Own_DLIS_Model.docx（Hao Zhang）', ZHAHAO_DOC_URL)

add_heading('8.2 构建 DLIS Service', 2)
add_para_with_link('参考：', 'How_to_Build_Your_Own_DLIS_Model.docx（Step 8）', ZHAHAO_DOC_URL, ' 创建 DLIS Service。')
add_para('操作步骤：', bold=True)
add_bullet('1. 打开 DLModelV2 - One Inference Portal，点击 New Model')
p = doc.paragraphs[-1]
# Replace bullet text with hyperlinked version
p.clear()
p.add_run('1. 打开 ')
add_hyperlink(p, 'DLModelV2 - One Inference Portal', ONE_INFERENCE_PORTAL)
p.add_run('，点击 New Model')
add_image('zhahao_image018.png', '图：One Inference Portal - 点击 New Model（来源：Hao Zhang 文档）')
add_bullet('2. 粘贴 Polaris Job Id，配置各页面参数：')
add_image('zhahao_image020.png', '图：配置 Key、Hardware、General、ACL 等页面（来源：Hao Zhang 文档）')
add_bullet('Key 页面：更新 Environment 和 Namespace')
add_bullet('Hardware 页面：选择部署目标机器')
add_bullet('General 页面：填写 DRI 联系人，设置 min/max instance 数量，Model Priority 设为 Test（生产环境则保持 Production）')
add_bullet('ACL 页面：添加访问控制 ACL')

add_heading('8.3 ACL 配置', 2)
add_para('DLIS Service 的 ACL 字符串包含多个证书指纹和 AAD 应用 ID，用于控制谁可以调用该服务：')
add_code('*:Certificate://Thumbprint/02AAAAA5AD...,*:AAD://appid/dda2a640-...,\n*:Certificate://Microsoft/dlis.si.advisoraggregator.trafficmanager.net,...')
add_note('ACL 配置错误会导致调用方收到 403 Forbidden。部署后如果请求被拒绝，优先检查 ACL 配置。')

add_bullet('3. 点击 VALIDATION 验证，成功后点击 SUBMIT 提交')
add_bullet('Prod 部署需要 bypass 流程（找人手动操作），比较卡时间（Siwen 经验）')

# ========== Section 9: 部署后验证与测试 ==========
add_heading('9. Step 7：部署后验证与测试', 1)
add_para('DLIS 正式部署成功后，需要发送请求验证服务是否正常工作。')

add_heading('9.1 Endpoint 命名规范', 2)
add_bullet('命名应该是描述性的、稳定的（如 PicassoAdsCreative.ZImage-V1）')
add_bullet('避免使用个人名字作为 endpoint 名称')
add_bullet('SI 和 Prod 不应共享同一个 endpoint（这是 pilot blocker）')

add_heading('9.2 请求 URL 格式', 2)
add_table(
    ['错误格式', '正确格式'],
    [
        ['http://WestUS2BE.bing.prod.dlis.binginternal.com:86/route/...', 'https://WestUS2.bing.prod.dlis.binginternal.com/route/...'],
        ['带 :8888 后缀', '不需要端口后缀'],
        ['/routebatch/', '/route/'],
    ]
)

add_heading('9.3 测试请求示例', 2)
add_code('''import requests

response = requests.post(
    "https://WestUS2.bing.prod.dlis.binginternal.com/route/PicassoAdsCreative.<ModelName>",
    cert=("private1.cer", "private1.key"),
    json={"prompt": "test input"},
    headers={"Content-Type": "application/json"},
    verify=False,
)
''')
add_para('需要客户端证书（.cer + .key 文件）进行认证，不能用 plain curl。')
add_link_note('证书类型与管理详见「第 11 节：SI/Prod 环境与证书管理」')
add_link_note('验证 Kusto 日志是否正常写入详见「第 12 节：Kusto 日志查看与调试」')

# ========== Section 10 ==========
add_heading('10. Step 8：Polaris Job 优化（可选）', 1)
add_note('此步骤为可选。如果跳过此步骤，需要联系 Fang Zhang 提交 bypass job。')
add_heading('10.1 量化', 2)
add_bullet('自动优化：通过 Polaris Job 内置的量化流程自动执行')
add_bullet('离线量化（推荐）：使用 AutoGPTQ、AutoAWQ 或 llm-compressor')
add_bullet('离线量化后上传量化模型到 Cosmos，vLLM 会从模型的 config.json 自动检测量化格式')

add_heading('10.2 其他优化', 2)
add_table(
    ['优化方式', '说明'],
    [
        ['Async API Call', '异步推理调用，提升吞吐量'],
        ['Continuous Batching', '连续批处理，显著提升吞吐量'],
    ]
)

# ================================================================
# PART 2: ENVIRONMENT & OPS
# ================================================================
doc.add_page_break()
p = doc.add_paragraph()
run = p.add_run('第二部分：环境与运维')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)

# ========== Section 11 ==========
add_heading('11. SI/Prod 环境与证书管理', 1)

add_heading('11.1 SI 与 Prod 环境分离', 2)
add_bullet('SI 环境用于测试验证，Prod 环境用于生产服务')
add_bullet('SI 和 Prod 共享同一个 endpoint 是 pilot blocker（Image Model Service 群组确认），正式上线必须分离')
add_bullet('Prod 部署需要 bypass 流程（找人手动操作），比较卡时间')

add_heading('11.2 证书类型与用途', 2)
add_table(
    ['证书类型', '格式', '用途', '位置'],
    [
        ['客户端认证证书', '.cer + .key', '调用 DLIS endpoint', '本地机器或服务器'],
        ['Kusto 日志证书', '.pfx', 'EventHub 写入 Kusto 日志', 'Cosmos /Model/ 目录'],
        ['SSL Keys', '.cert + .key', '服务端 SSL', '服务器特定路径'],
    ]
)

add_para('证书文件位置参考（Siwen 提供）：')
add_bullet('SSL keys: 10.224.120.197 /home/siwen/relevance/deploy')
add_bullet('新证书位于 Cosmos: cosmos09.osdinfra.net/.../ImgLPRelevance6/')

add_heading('11.3 证书过期管理', 2)
add_bullet('2026 年 4 月底证书到期，新证书已改为 .pfx 格式')
add_bullet('证书过期后服务将无法认证，需要提前更新')
add_bullet('建议设置证书到期提醒，提前 2 周更新')

# ========== Section 12 ==========
add_heading('12. Kusto 日志查看与调试', 1)

add_heading('12.1 Kusto 日志环境路由', 2)
add_para('关键发现（Siwen 确认）：证书决定了你能看到哪个环境的日志。', bold=True)
add_bullet('使用 SI 证书 → 只能看到 SI 环境的日志')
add_bullet('使用 Prod 证书 → 只能看到 Prod 环境的日志')
add_bullet('如果 namespace 用的 prod 证书用的 si，日志会写到错误的库')
add_note('曾发现 SI Kusto 日志错误地写入了 Prod DB（Image Model Service 群组），需注意配置一致性。')

add_heading('12.2 Kusto 查询', 2)
add_bullet('SI 环境：bingadsppe.AdInsightMT')
add_bullet('Prod 环境：bingads.BingAdsTracing')
add_table_with_links(
    ['环境', '链接'],
    [
        ['PROD Kusto', ('https://bingads.kusto.windows.net/', KUSTO_PROD_URL)],
        ['SI Kusto', ('https://bingadsppe.kusto.windows.net/', KUSTO_SI_URL)],
        ['DLIS Jarvis Dashboard (Prod)', ('DLIS Model Metrics | Jarvis', JARVIS_PROD_URL)],
        ['DLIS Jarvis Dashboard (SI)', ('DLIS Model Metrics | Jarvis (SI)', JARVIS_SI_URL)],
    ]
)
add_doc_link('Auto Image Service DLIS Documentation（ChunChen）— 含 Kusto 日志分析示例', CHUNCHEN_DOC_URL)
add_code('''appsvc_info | union appsvc_warn | union appsvc_err
| where Timestamp > ago(30min)
| where ApplicationName == 'ImgLPRelevanceModel'
''')
add_para('ApplicationName 的值来自 config 文件中 application_name 字段（Siwen 确认）。')

add_heading('12.3 Polaris Log', 2)
add_bullet('Polaris log 可以看到服务启动日志，但没有容器内的 print log（Siwen 确认）')
add_bullet('如果需要更详细的日志，用 Kusto 或 Central Log')
p = doc.paragraphs[-1]
p.clear()
p.add_run('如果需要更详细的日志，用 Kusto 或 ')
add_hyperlink(p, 'Central Log', DLIS_WIKI_CENTRAL_LOG)
add_link_note('Central Log 查询详见「附录 E」')

add_heading('12.4 本地测试 Kusto Log', 2)
add_para('重要原则：在本地开发阶段就应集成 Kusto log 输出。不要等到部署到线上才开始接入 Kusto 日志——本地 Docker 测试时就应该能看到 Kusto log，这样可以提前发现日志格式、认证配置、EventHub 连接等问题。', bold=True)
add_para('本地测试只需要：① 正确的 PFX 证书文件；② 对应环境的 EventHub namespace 配置；③ 网络能访问 EventHub 端点。满足这三个条件，本地 Docker 容器中的日志就会实时发送到 Kusto，可以直接在 Kusto Explorer 中查询验证。')
add_para('使用 AAD Bearer token（MSAL + PFX 证书）认证的测试脚本：')
add_code('''import msal, requests
from cryptography.hazmat.primitives.serialization.pkcs12 import load_key_and_certificates

# 1. 加载 PFX 证书 → 获取 private_key, thumbprint, public_certificate
# 2. 获取 AAD Bearer token
app = msal.ConfidentialClientApplication(
    client_id=CLIENT_ID,
    authority=f"https://login.microsoftonline.com/{TENANT_ID}",
    client_credential={"private_key": ..., "thumbprint": ..., "public_certificate": ...},
)
result = app.acquire_token_for_client(scopes=[DLIS_SCOPE])

# 3. 发送带 tracking_data 的请求
payload = {
    "prompt": "test", "width": 1344, "height": 768,
    "tracking_data": {
        "requestid": f"test-{uuid.uuid4().hex[:8]}",
        "trackingid": f"test-{uuid.uuid4().hex[:8]}",
        "callername": "local_test",
    },
}
resp = requests.post(API_URL, json=payload,
                     headers={"Authorization": f"Bearer {token}"},
                     verify=False, timeout=120)
''')

add_heading('12.5 Kusto 日志常见问题与调试经验', 2)
add_para('以下经验总结自 ZImage、Gemma4 等模型部署过程中的实际调试记录和团队讨论。')

add_heading('问题 1：EventHub 认证错误被静默吞掉', 3)
add_para('现象：本地测试有正常返回，但 Kusto 中始终看不到任何日志。')
add_para('根因：kusto_log.py 中 EventHub 发送失败时使用了 try/except 捕获所有异常并 pass，导致证书或 namespace 配置错误完全无感知。')
add_code('''# ❌ 错误写法 — 静默吞掉认证失败
try:
    client.send(event_data)
except Exception:
    pass  # 日志丢失，无任何提示

# ✅ 正确写法 — fail fast，暴露配置问题
try:
    client.send(event_data)
except Exception as e:
    logger.error(f"EventHub send failed: {e}", exc_info=True)
    raise  # 首次失败应立即暴露''')

add_heading('问题 2：record.msg vs record.getMessage()', 3)
add_para('现象：Kusto 日志中的消息缺少变量替换，显示原始模板字符串（如 "%s loaded in %d seconds"）。')
add_para('根因：kusto_log.py 中使用 record.msg 获取日志消息，但 record.msg 是未格式化的模板。应使用 record.getMessage() 获取完整格式化后的字符串。')
add_code('''# ❌ record.msg → "Model %s loaded in %d seconds"
# ✅ record.getMessage() → "Model gemma4 loaded in 42 seconds"''')

add_heading('问题 3：进程 crash 时 Kusto 日志丢失', 3)
add_para('现象：模型加载阶段 OOM/CUDA 错误导致进程崩溃，但 Kusto 中没有任何错误日志。')
add_para('根因：KustoHandler 使用 ScheduledBatchSender 定时批量发送。进程崩溃时 scheduler 线程随之终止，缓冲区中的日志全部丢失。')
add_code('''# 解决方案：在 crash handler 中手动 flush
import signal, atexit

def flush_kusto_on_exit():
    for handler in logging.root.handlers:
        if hasattr(handler, 'flush'):
            handler.flush()

atexit.register(flush_kusto_on_exit)
signal.signal(signal.SIGTERM, lambda *_: (flush_kusto_on_exit(), sys.exit(1)))''')

add_heading('问题 4：SI/Prod 证书与 namespace 不匹配', 3)
add_para('现象：日志发送无报错（如果异常被捕获），但 Kusto 查不到任何记录。')
add_para('根因：使用 SI 环境的证书连接 Prod 的 EventHub namespace（或反之），EventHub 鉴权通过但消息被路由到错误的环境。')
add_bullet('SI 环境：namespace 带 "si" 后缀，使用 SI 证书')
add_bullet('Prod 环境：namespace 不带后缀，使用 Prod 证书')
add_para('建议：将环境配置放在 settings.json 中，通过环境变量切换，避免硬编码。', bold=True)

add_heading('问题 5：Logger 级别默认 WARNING', 3)
add_para('现象：代码中有 logger.info() 调用，但 Kusto 中只能看到 WARNING 及以上级别的日志。')
add_para('根因：Python logging 的子 logger 默认继承 root logger 的级别（WARNING）。如果没有显式设置，所有 INFO 和 DEBUG 日志都会被过滤。')
add_code('''# 必须显式设置 logger 级别
logger = logging.getLogger("dlis_model")
logger.setLevel(logging.INFO)  # 不设置则默认继承 WARNING''')

add_heading('问题 6：本地测试就应该验证 Kusto 日志', 3)
add_para('重要：不要把 Kusto 日志验证留到线上部署阶段。在本地 Docker 测试时就应该集成 Kusto log，确保日志能正常发送到 EventHub。')
add_para('本地测试 Kusto 日志的前提条件：')
add_bullet('有正确环境的 PFX 证书文件（SI 或 Prod）')
add_bullet('settings.json 中配置了正确的 EventHub namespace 和 topic')
add_bullet('网络能访问 EventHub 端点（公司网络或 VPN）')
add_para('验证步骤：')
add_bullet('本地 Docker 启动后，发送测试请求，观察容器日志中是否有 EventHub 发送成功/失败的输出')
add_bullet('同时在 Kusto Explorer 中查询对应环境的表，确认日志已到达（通常延迟 1-2 分钟）')
add_bullet('如果网络受限无法连接 EventHub，可临时使用 console handler 验证日志格式，但应尽早在能连接 EventHub 的环境中做完整验证')

add_heading('问题 7：EventHub 四个 Topic 的用途', 3)
add_para('DLIS EventHub 提供四个 topic，需要根据日志类型发送到对应的 topic：')
add_bullet('appsvc_info — 常规信息日志（模型加载、请求处理等）')
add_bullet('appsvc_warn — 警告日志（非致命错误、性能降级等）')
add_bullet('appsvc_err — 错误日志（异常、crash 信息等）')
add_bullet('appsvc_perf — 性能日志（推理延迟、吞吐量等指标）')
add_para('注意：如果只往 appsvc_info 发送，在 Kusto 中查询 errors 表将看不到任何结果。', bold=True)

# ========== Section 13 ==========
add_heading('13. 常见问题与解决方案', 1)

problems = [
    ('问题 1：容器 OOM Killed（CPU 内存，非 GPU）',
     'OaasWrapper 找不到 _opt 目录 → fallback 到 BaseLLM → CPU 加载全部权重 → 超过 CPU 内存限制',
     '确保 _opt 目录结构正确（方案 A），或直接使用 vllm.LLM()（方案 B，推荐）'),
    ('问题 2：CUDA_VISIBLE_DEVICES UUID 格式',
     'DLIS 设置 CUDA_VISIBLE_DEVICES 为 GPU UUID 格式，vLLM 内部 int() 转换失败',
     'run.sh 中 unset CUDA_VISIBLE_DEVICES + model.py 中做 UUID → 整数索引转换'),
    ('问题 3：/Model 是只读文件系统',
     'DLIS 将 Cosmos 数据以只读方式挂载到 /Model',
     'Writable Mirror 方案：在 /tmp 创建镜像目录，用 symlink 指向只读模型文件'),
    ('问题 4：Unable to find exposed port 8888',
     '通常不是端口问题，而是模型加载 crash（OOM 等），HTTP server 从未启动',
     '先解决模型加载问题，同时确保 Dockerfile 有 EXPOSE 8888'),
    ('问题 5：Pipeline 构建 pip install 超时',
     'CI agent 无法直连 pypi.org',
     '添加清华 PyPI 镜像（见附录 B）'),
    ('问题 6：OaasWrapper _create_runner() 静默吞异常',
     '_create_runner() 有 try/except 捕获所有异常并 print()，然后 return None',
     '修改为 logger.error(exc_info=True)，或直接使用方案 B'),
    ('问题 7：opt_type.txt 换行符',
     'echo "llm" 写入时带了末尾换行符',
     '用 printf "llm" 代替 echo "llm"'),
    ('问题 8：Cosmos 挂载只挂载一级目录',
     '把证书放在子文件夹里面，提示找不到',
     '证书等文件放在 Cosmos 根目录，不要嵌套在子文件夹中'),
    ('问题 9：log 里用的 docker 和 config 配置的不一致',
     '配置文件与实际运行的镜像不一致',
     '检查 Polaris Job 配置中的 ModelPath 是否指向正确的镜像 tag'),
    ('问题 10：本地测试有返回但 Kusto 日志为空',
     '使用了错误环境的证书（如 namespace 用 prod，证书用 si）',
     '确保 namespace 和证书环境一致（见第 11、12 节）'),
]

for title, cause, fix in problems:
    add_heading(title, 2)
    add_para(f'根因：{cause}')
    add_para(f'解决：{fix}')

# ================================================================
# APPENDICES
# ================================================================
doc.add_page_break()
p = doc.add_paragraph()
run = p.add_run('附录')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)

# ========== Appendix A ==========
add_heading('附录 A：Base Docker 镜像选择与更新', 1)

add_heading('A.1 快速构建方案（推荐）：Dockerfile_vllm_fast', 2)
add_code('''FROM vllm/vllm-openai:latest
# 已包含 vllm、torch、transformers，无需编译
# 如需特定 transformers 版本：
RUN python3 -m pip install transformers==5.5.3
''')
add_bullet('优点：构建 < 1 秒，包含预编译的 vllm + torch')
add_bullet('缺点：latest 版本不可控，可能被上游更新')

add_heading('A.2 完整构建方案：Dockerfile_vllm_0.10.0', 2)
add_code('''FROM nvidia/cuda:12.8.1-devel-ubuntu22.04

# 关键：先装正确版本的 torch，再装 vllm
RUN pip install torch==2.10.0 torchvision==0.25.0 \\
    --index-url https://download.pytorch.org/whl/cu128
RUN pip install vllm==0.19.0
''')

add_heading('A.3 版本兼容性经验', 2)
add_table(
    ['问题', '根因', '解决'],
    [
        ['vllm/_C.abi3.so: undefined symbol', 'torch ABI 不匹配', '先装 torch，requirements 不加 torch/torchvision'],
        ['torchvision::nms does not exist', '版本不匹配', '确保 torchvision 与 torch 对应'],
        ['num_scheduler_steps 不被接受', 'vllm 0.19.0 移除', '从 vllm_runner.py 中删除'],
        ['Gemma4VideoProcessor requires Torchvision', '卸载 torchvision', '装回正确版本'],
    ]
)
add_para('原则：requirements-vllm.txt 不要加 torch、torchvision、transformers，由 Dockerfile 统一管理。', bold=True)

add_heading('A.4 本地手动构建测试', 2)
add_code('''cd OaaS_LLMTemplate
IMAGE_TAG="local-test"

# Block 1: 构建基础镜像
sudo docker build -t my-vllm-base:$IMAGE_TAG \\
    --file pipeline/Dockerfile_vllm_fast pipeline/

# Block 2: 安装 OaaS 代码和依赖
sudo docker run -d --name temp my-vllm-base:$IMAGE_TAG sleep infinity
sudo docker cp . temp:/
sudo docker exec temp chmod +x /dlis_model/run.sh /dlis_model/async_run.sh /LLMModelOptimization.sh
sudo docker exec temp python3 -m pip install -r /requirements-common.txt
sudo docker exec temp python3 -m pip install -r /requirements-vllm.txt
sudo docker exec temp python3 -m pip install -e /
sudo docker commit temp my-final:$IMAGE_TAG
sudo docker rm -f temp
''')

# ========== Appendix B ==========
add_heading('附录 B：清华镜像源配置', 1)
add_para('DLIS CI pipeline agent 和 Docker build 可能无法直连 pypi.org。')
add_heading('B.1 CI Pipeline（azure-pipelines-unified.yml）', 2)
add_code('''variables:
  PIP_INDEX_URL: https://pypi.tuna.tsinghua.edu.cn/simple
  UV_INDEX_URL: https://pypi.tuna.tsinghua.edu.cn/simple
''')
add_heading('B.2 Docker Build（build_vllm_image.sh）', 2)
add_code('''docker build \\
  --build-arg PIP_INDEX_URL=https://pypi.tuna.tsinghua.edu.cn/simple \\
  --build-arg UV_INDEX_URL=https://pypi.tuna.tsinghua.edu.cn/simple \\
  --build-arg PIP_EXTRA_INDEX_URL=https://pypi.org/simple \\
  ...

# docker exec 安装阶段：
PIP_ARGS="-i https://pypi.tuna.tsinghua.edu.cn/simple --extra-index-url https://pypi.org/simple"
docker exec container pip install $PIP_ARGS -r requirements.txt
''')
add_heading('B.3 Dockerfile 接收参数', 2)
add_code('''ARG PIP_INDEX_URL
ARG UV_INDEX_URL
ARG PIP_EXTRA_INDEX_URL
ARG UV_EXTRA_INDEX_URL
''')
add_note('注意：不能用 docker exec -e 环境变量方式传清华源参数，必须用 --build-arg 或命令行参数。')

# ========== Appendix C ==========
add_heading('附录 C：Model 代码编写指南', 1)

add_heading('C.1 model.py 核心结构', 2)
add_code('''class ModelImp:
    def __init__(self):
        # 模型初始化：加载引擎、配置参数
        pass

    def Eval(self, data):
        # 单条推理
        pass

    def EvalBatch(self, data_list):
        # 批量推理
        pass
''')

add_heading('C.2 方案 A：使用 OaasWrapper（适合简单 LLM 推理）', 2)
add_code('''from llm_opt.oaas_wrapper_v2 import OaasWrapper

class ModelImp:
    def __init__(self):
        self.oaas_wrapper = OaasWrapper("model_dir_name", is_llm_model=True)

    def Eval(self, data):
        prompts = preprocess(data)
        outputs = self.oaas_wrapper.run(prompts)
        return postprocess(outputs)
''')
add_bullet('需要 _opt 目录下的 opt_type.txt 和 best_setting.json')
add_note('注意：best_setting.json 中不要设 quantization 字段。vLLM 会从模型 config.json 自动检测。')

add_heading('C.3 方案 B：直接使用 vLLM（推荐，更可控）', 2)
add_code('''from vllm import LLM, SamplingParams

class ModelImp:
    def __init__(self):
        # CUDA_VISIBLE_DEVICES UUID 修复
        cuda_env = os.environ.get("CUDA_VISIBLE_DEVICES", "")
        if cuda_env and not cuda_env.replace(",", "").isdigit():
            gpu_count = len(cuda_env.split(","))
            os.environ["CUDA_VISIBLE_DEVICES"] = ",".join(str(i) for i in range(gpu_count))

        self.llm = LLM(
            model="/Model/model-name",
            tensor_parallel_size=1,
            gpu_memory_utilization=float(os.environ.get("GPU_MEMORY_UTILIZATION", "0.9")),
            trust_remote_code=True, dtype="auto", max_model_len=8192,
            enable_prefix_caching=True,
        )

        self.sampling_params = SamplingParams(
            temperature=1.0, top_p=0.95, top_k=64,
            max_tokens=128, stop=["</end_token>"],
        )

    def Eval(self, data):
        prompts = preprocess(data)
        outputs = self.llm.generate(prompts, self.sampling_params)
        return [o.outputs[0].text for o in outputs]
''')
add_para('方案 B 的优势：', bold=True)
add_bullet('去掉 OaasWrapper 中间层，初始化失败直接报错')
add_bullet('不需要 _opt 目录和 best_setting.json')
add_bullet('推理参数直接写在代码中，透明可控')
add_note('Siwen 也采用直接使用 vllm.LLM 而非 OaasWrapper 的方案，这已成为团队共识。')

add_heading('C.4 非 LLM 模型（如 ZImage diffusion）', 2)
add_code('''class ModelImp:
    def __init__(self):
        from diffusers import ZImagePipeline
        model_path = os.environ.get("ZIMAGE_MODEL_PATH", "/Model/model_name")
        self.pipe = ZImagePipeline.from_pretrained(model_path, torch_dtype=torch.bfloat16)
        self.pipe.to("cuda:0")

    def Eval(self, data):
        result = self.pipe(prompt=data["prompt"], width=data["width"], height=data["height"])
        ...
''')

add_heading('C.5 多步推理（如 Gemma4 Two-Step）', 2)
add_code('''def Eval(self, data):
    # Step 1：生成 scene concepts
    step1_outputs = self._run_vllm(step1_prompts, self.step1_params)

    # Step 2：展开为详细 prompts
    step2_outputs = self._run_vllm(step2_prompts, self.step2_params)
    return self.processor.postprocess(step2_outputs)
''')
add_note('关键：Step 1 和 Step 2 应使用独立的 SamplingParams（不同的 max_tokens 和 stop tokens）。')

add_heading('C.6 Tokenizer Thinking Mode 问题', 2)
add_para('某些量化模型的 tokenizer chat_template 内置了 thinking prefix，需要 patch：')
add_code('''tokenizer = self.llm.get_tokenizer()
if hasattr(tokenizer, 'chat_template') and '<|channel>thought' in (tokenizer.chat_template or ''):
    tokenizer.chat_template = tokenizer.chat_template.replace(
        "<|channel>thought\\n<channel|>", ""
    )
''')

# ========== Appendix D ==========
add_heading('附录 D：外部 settings.json 配置覆盖', 1)

add_heading('D.1 背景', 2)
add_para('config.py 中的配置如果硬编码，切换 si/prod 环境需要改代码重新构建镜像。参考 Hanbang（user/hanbangliang/img-outpainting-v1 分支）的实现，通过 Cosmos 上的外部 JSON 文件覆盖。')

add_heading('D.2 实现方式', 2)
add_code('''from pydantic_settings import BaseSettings, SettingsConfigDict

DEFAULT_SETTINGS_JSON_PATH = os.environ.get("SETTINGS_JSON_PATH", "/Model/settings.json")

class Settings(BaseSettings):
    model_config = SettingsConfigDict(extra="ignore")
    eventhub_namespace: str = "aggregation-logging.servicebus.windows.net"
    certificate_path: str = os.path.join("/Model", "AggSvcAuthCert-prod.pfx")
    ...
''')

add_heading('D.3 配置优先级', 2)
add_bullet('1. init kwargs（代码显式传参）')
add_bullet('2. 环境变量（如 EVENTHUB_NAMESPACE=...）')
add_bullet('3. JSON 文件（/Model/settings.json）')
add_bullet('4. 代码默认值')

add_heading('D.4 使用方式', 2)
add_para('在 Cosmos 根目录放置 settings.json，只写需要覆盖的字段：')
add_code('''{"eventhub_namespace": "aggregation-si-logging.servicebus.windows.net",
 "certificate_path": "/Model/AggSvcAuthCert-si.pfx"}
''')
add_para('不放 settings.json 则使用代码默认值（prod 环境）。')

add_heading('D.5 EventHub Credential 容错', 2)
add_code('''def _try_get_credential(tenant_id):
    try:
        return CertificateCredential(...)
    except Exception as e:
        logger.info("EventHub credential unavailable, kusto logs local-only")
        return None
''')
add_para('证书加载失败不应阻塞容器启动，静默降级为本地日志。')

# ========== Appendix E ==========
add_heading('附录 E：Central Log 调试工具', 1)
add_para('使用 Central Log 调试 Polaris Job，查看容器输出日志（Desheng 推荐）：')
add_code('''SELECT machine_name, log_level, log_time, description
FROM dlissensitivelog
WHERE file_name LIKE 'DLMSUserLog_ContainerOutput%.log'
  AND log_time BETWEEN TIMESTAMP '2026-02-27 00:00:00'
                   AND TIMESTAMP '2026-02-28 18:00:00'
  AND machine_name = '<your_machine_name>'
LIMIT 10000;
''')

# ========== Appendix F ==========
add_heading('附录 F：参考链接', 1)
add_table_with_links(
    ['资源', '链接/说明'],
    [
        ['DLIS LLM Engine 官方文档', ('Optimizing a LLM model with DLIS LLM engine', DLIS_WIKI_LLM_ENGINE)],
        ['OaaS Template 仓库', ('OaaS_LLMTemplate（ADO Git）', OAAS_LLM_TEMPLATE_REPO)],
        ['Gen1→Gen2 数据迁移工具', ('Data Transfer Tools（DLIS Wiki）', DLIS_WIKI_DATA_TRANSFER)],
        ['Central Log 查询指南', ('How to Use Central Log（DLIS Wiki）', DLIS_WIKI_CENTRAL_LOG)],
        ['DLIS 模型构建指南（Hao Zhang）', ('How_to_Build_Your_Own_DLIS_Model.docx', ZHAHAO_DOC_URL)],
        ['OaaS 部署指南 v2（ChangXu）', ('DLIS_Model_DeploymentWith_OaaS_v2.docx', XUCHA_DOC_URL)],
        ['Kusto 日志（SI）', ('bingadsppe.AdInsightMT | Azure Data Explorer', KUSTO_SI_URL)],
        ['Kusto 日志（Prod）', ('bingads.BingAdsTracing | Azure Data Explorer', KUSTO_PROD_URL)],
        ['Auto Image Service DLIS 文档（ChunChen）', ('Auto Image Service — DLIS Model Documentation', CHUNCHEN_DOC_URL)],
        ['DLIS 部署流程讲解（Siwen & Desheng）', ('Call with Desheng Cui — 2026-03-26 录屏', SIWEN_DESHENG_RECORDING)],
    ]
)
add_doc_link('DLIS_Model_DeploymentWith_OaaS_v2.docx（ChangXu）', XUCHA_DOC_URL)
add_doc_link('How_to_Build_Your_Own_DLIS_Model.docx（Hao Zhang）', ZHAHAO_DOC_URL)
add_doc_link('Auto Image Service — DLIS Model Documentation（ChunChen）', CHUNCHEN_DOC_URL)
add_doc_link('Call with Desheng Cui — 2026-03-26 DLIS 部署流程录屏（Siwen & Desheng）', SIWEN_DESHENG_RECORDING)

# ========== Appendix G ==========
add_heading('附录 G：Gemma4 DLIS 部署迭代总结', 1)
add_table(
    ['部署 #', '问题', '修复'],
    [
        ['#1-#3', '_opt 目录不可见 → OOM', '确认 Cosmos 同步无法解决'],
        ['#4', '/Model 只读无法创建 _opt', 'Writable mirror 方案'],
        ['#5', 'vLLM 初始化失败被静默吞掉', '_create_runner() 改为抛异常'],
        ['#6', '改为 raise 后 crash loop', '改回 return None + logger.error'],
        ['#7', 'CUDA_VISIBLE_DEVICES UUID 格式', 'UUID → 整数索引转换'],
        ['#8', 'Cosmos 目录 + 旧 model.py 残留', '新建扁平目录 + 去掉 OaasWrapper'],
    ]
)
add_para('核心教训：OaasWrapper 中间层复杂度远大于收益。推荐直接使用 vllm.LLM()（方案 B）。', bold=True)

# ========== Appendix H ==========
add_heading('附录 H：修订历史', 1)
add_table(
    ['版本', '日期', '作者', '说明'],
    [
        ['v2', '2026-02-26', 'ChangXu', '初始文档，完整部署指南'],
        ['v4', '2026-04-22', 'Jinjin Chen', '整合 Gemma4/ZImage 实战经验'],
        ['v5', '2026-04-22', 'Jinjin Chen', '整合 ChangXu v2、settings.json、Kusto 测试'],
        ['v5.1', '2026-04-22', 'Jinjin Chen', '重构文档结构：主流程 + 附录；整合团队 Teams 聊天经验'],
    ]
)

# Save
output_path = r"C:\Users\jinjinchen\OneDrive - Microsoft\DLIS_Model_Deployment_Guide_v5_3.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
